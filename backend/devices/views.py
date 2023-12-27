from abc import ABC
from dataclasses import dataclass
from pickle import TRUE
from sqlite3 import connect

# from turtle import window_width
from django.db.models.query import QuerySet

# from html5lib import serialize
from rest_framework.parsers import JSONParser
from rest_framework.decorators import parser_classes
from rest_framework.decorators import api_view
from rest_framework.decorators import api_view, parser_classes
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from django.db.models.query_utils import Q
from django.db.models.signals import post_save
from django.http import Http404, response
from requests.api import request
from rest_framework import serializers, views
from django.core import serializers
from sshutil import SshClient
from devices import tasks
from django.http import HttpResponse
from rest_framework import status

# from devices.scanner import arp_scan
from rest_framework.response import Response
from rest_framework.views import APIView
from logging import log, raiseExceptions
from rest_framework import permissions, generics
from devices.serializers import *  # AddProcessListWithIPSerializer, DevicesSerializer, IntegrityCheckSerializer, ProcessHashSerializer, ProcessListSerializer, SyscallListSerializer, UpdateDeviceSerializer, AlertsSerializer, UpdateIPFrequencySerializer, UpdateDeviceTelnetPasswordSerializer, UpLoadImageSerializer, FileExampleSerializer   , ModelMachineLearningSerializer,IpsTrackingSerializer, AddCountIpsTrackingSerializer, OneSignalSerializer
from devices.models import *  # Devices, Alerts, IntegrityCheck, ProcessHash, ProcessList, SyscallList, SyscallListSerializerCustom, save_syscall , ModelML, IpsTracking,OneSignal
from django.shortcuts import render
from rest_framework import viewsets
from drf_spectacular.utils import extend_schema, OpenApiParameter
from drf_spectacular.types import OpenApiTypes
from rest_framework.decorators import action
from rest_framework.pagination import PageNumberPagination
from channels.layers import get_channel_layer
from asgiref.sync import async_to_sync
from django.core.paginator import Paginator
from django.db.models.signals import post_save
from .tools.create_info_file import CreateInfoFile
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework import filters
from rest_framework.parsers import (
    FormParser,
    MultiPartParser,
    JSONParser,
    FileUploadParser,
)
from django.core import serializers
import os
import logging
import subprocess
import json
import string
import random

logger = logging.getLogger(__name__)
import re
import subprocess
import requests
from packaging import version
import io
from django.http import FileResponse
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import xlsxwriter
from io import BytesIO
from django.http import HttpResponse
from pathlib import Path

# Create your views here.
from auth_user.models import AuthUser
from onesignal_sdk.client import Client
from onesignal_sdk.error import OneSignalHTTPError
from .config import APP_ID, REST_API_KEY, USER_AUTH_KEY
import psutil
import os
from os.path import isfile
from datetime import datetime, timedelta, timezone
from docx import Document
import aspose.words as aw
from elasticsearch import Elasticsearch
from dotenv import dotenv_values
from django.shortcuts import get_object_or_404
from .tools import validateip, capture, SftpRequest
from . import extract_pcap
from .permissions import HasAPIKey
from IoTAnalyzer.env_dev import *
import random
import redis
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
from .extract_files_frompcap import extract_allFolder
from django.http import HttpResponse
from IoTAnalyzer.env_dev import *
import threading
import time
import multiprocessing

APP_ID = "c2234356-b7c0-4ba4-b786-daecff964a82"
REST_API_KEY = "NjI3NTJkNmItYjE5My00M2NmLWJhM2MtN2RhNzFkM2JiYzNl"
USER_AUTH_KEY = "NWIxN2QzN2YtYTQ2Mi00ZjllLTk5OWMtYTYwMDk3MWQyNTAw"
pathEnv = "/backend/.env.dev"


def get_headers():
    url = f'{os.getenv("CENTER_SERVER_IP", "http://192.168.10.162:9000")}/auth/token/login/'
    username = os.environ.get("CENTER_USERNAME", "tomcat")
    password = os.environ.get("CENTER_PASSWORD", "tomcat")
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    x = requests.post(
        url,
        data=json.dumps({"username": username, "password": password}),
        headers=headers,
    )
    content = json.loads(x.content)

    auth_token = content["auth_token"]
    headers = {
        "Authorization": f"Token {auth_token}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    return headers


def action_device_norm(type, data):
    r = redis.Redis(host="redis-agent", port=6379, db=0)
    if type == "create" or type == "update":
        template_data = {
            "net_rxkBs": {
                "min_a": data["net_rxkBs_min"],
                "max_a": data["net_rxkBs_max"],
            },
            "net_txkBs": {
                "min_a": data["net_txkBs_min"],
                "max_a": data["net_txkBs_max"],
            },
            "net_tpcks": {
                "min_a": data["net_tpcks_min"],
                "max_a": data["net_tpcks_max"],
            },
            "net_rpcks": {
                "min_a": data["net_rpcks_min"],
                "max_a": data["net_rpcks_max"],
            },
            "io_rkbs": {"min_a": data["io_rkbs_min"], "max_a": data["io_rkbs_max"]},
            "io_wkbs": {"min_a": data["io_wkbs_min"], "max_a": data["io_wkbs_max"]},
        }
        if r.get("device_norm"):
            data_device_norm = json.loads(r.get("device_norm").decode())
        else:
            data_device_norm = {}
        data_device_norm[str(data["device"])] = template_data
        r.set("device_norm", json.dumps(data_device_norm))
    else:
        if r.get("device_norm"):
            data_device_norm = json.loads(r.get("device_norm").decode())
        else:
            data_device_norm = {}
        if str(data["device"]) in data_device_norm:
            del data_device_norm[str(data["device"])]
        r.set("device_norm", json.dumps(data_device_norm))


def action_check_device_norm(data, status=False, trace_pcap=False):
    r = redis.Redis(host="redis-agent", port=6379, db=0)
    template_data = {
        "status": status,
        "time": int(datetime.now().timestamp()),
        "trace_pcap": trace_pcap,
    }
    if r.get("check_device_norm"):
        data_check_device_norm = json.loads(r.get("check_device_norm").decode())
    else:
        data_check_device_norm = {}
    data_check_device_norm[str(data["device"])] = template_data
    r.set("check_device_norm", json.dumps(data_check_device_norm))
    return data_check_device_norm


class AutoChangePw(APIView):
    def get(self, request):
        contents = Path(pathEnv).read_text()
        arr = contents.split(chr(10))
        autoChangePw = ""
        for item in arr:
            if item.find("AUTO_CHANGE_PWD_SECONDS") != -1:
                autoChangePw = item
        value = autoChangePw.split("=")
        return HttpResponse(value[1])


class PutAutoChangePw(APIView):
    def put(self, request, seconds=(60 * 60 * 24)):
        contents = Path(pathEnv).read_text()
        arr = contents.split(chr(10))
        for i in range(len(arr)):
            if arr[i].find("AUTO_CHANGE_PWD_SECONDS") != -1:
                arr[i] = "AUTO_CHANGE_PWD_SECONDS=" + str(seconds)
        res = ""
        for x in arr:
            res += x + chr(10)
        open(pathEnv, "w").close()
        f = open(pathEnv, "w")
        f.write(res)
        f.close()
        return HttpResponse("Changed days auto change password.")


class ExportXLS(APIView):
    def get(self, request):
        # create our spreadsheet.  I will create it in memory with a StringIO
        try:
            xxx_data = json.loads(request.query_params.get("data", None))
            result = Alerts.objects.all()
            if request.query_params.get("ip", None) != None:
                result = result.filter(ip__icontains=request.query_params["ip"])
            if request.query_params.get("message", None) != None:
                result = result.filter(
                    message__icontains=request.query_params["message"]
                )
            if request.query_params.get("address", None) != None:
                result = result.filter(
                    address__icontains=request.query_params["address"]
                )

            if xxx_data.get("timestamp", None) != None:
                # aaa = datetime.strptime(xxx_data.get('timestamp', None)[0], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                # bbb = datetime.strptime(xxx_data.get('timestamp', None)[1], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                result = result.filter(
                    timestamp__range=(
                        xxx_data.get("timestamp", None)[0],
                        xxx_data.get("timestamp", None)[1],
                    )
                )
            i = 0
            xyz = result
            if request.query_params["all_device"] == "false":
                if xxx_data.get("device", None) != None:
                    for id in xxx_data.get("device", None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0

            if xxx_data.get("type", None) != None:
                for id in xxx_data.get("type", None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = abc | result.filter(type=id)
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get("status", None) != None:
                for id in xxx_data.get("status", None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            worksheet = workbook.add_worksheet()
            worksheet.write("A1", "STT")
            worksheet.write("B1", "Mã thiết bị")
            worksheet.write("C1", "Địa chỉ ip")
            worksheet.write("D1", "Cảnh báo")
            worksheet.write("E1", "Hash")
            worksheet.write("F1", "PID")
            worksheet.write("G1", "Thời gian")
            worksheet.write("H1", "Địa chỉ")
            for i in range(len(xyz)):
                worksheet.write(f"A{i + 2}", i + 1)
                worksheet.write(f"B{i + 2}", xyz[i].device_id)
                worksheet.write(f"C{i + 2}", xyz[i].ip)
                worksheet.write(f"D{i + 2}", xyz[i].message)
                worksheet.write(f"E{i + 2}", xyz[i].hash)
                worksheet.write(f"F{i + 2}", xyz[i].pid)
                worksheet.write(f"G{i + 2}", str(xyz[i].timestamp))
                worksheet.write(f"H{i + 2}", str(xyz[i].address))
            workbook.close()

            # create a response
            response = HttpResponse(content_type="application/vnd.ms-excel")

            # tell the browser what the file is named
            response["Content-Disposition"] = 'attachment;filename="alerts.xlsx"'

            # put the spreadsheet data into the response
            response.write(output.getvalue())
            return response
        except Exception as e:
            logging.error(e)
            return Response(
                data={"status": False, "Log bug": str(e)},
                status=status.HTTP_400_BAD_REQUEST,
            )


class GetExportXLSURL(APIView):
    def get(self, request):
        url = request.get_full_path()
        url = url[: url.find("/url")] + url[url.find("/url") + 4 :]
        return Response(url)


class update_rule(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, format=None):
        tasks.publish_message({"type": "update_rule"})
        return Response({"status": "doing task..."})


class kill_agent_by_id(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get_object(self, pk):
        try:
            return Devices.objects.get(pk=pk)
        except Devices.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        device = self.get_object(pk)
        serializer = DevicesSerializer(device, context={"request": request})
        r = redis.Redis(host="redis-agent", port=6379, db=0)
        # logger.info(serializer.data)
        device_id = serializer.data["id"]
        time_check = datetime.now().timestamp()
        if r.get("device_check_have_data"):
            device_check_have_data = json.loads(r.get("device_check_have_data"))
            if device_check_have_data.get(f"{device_id}"):
                device_check_have_data[f"{device_id}"]["status"] = "GOTACTU"
                device_check_have_data[f"{device_id}"]["timestamp"] = time_check
            else:
                device_check_have_data[f"{device_id}"] = {
                    f"{device_id}": {"status": "GOTACTU", "timestamp": time_check}
                }
            r.set("device_check_have_data", json.dumps(device_check_have_data))
        else:
            device_check_have_data = {
                f"{device_id}": {"status": "GOTACTU", "timestamp": time_check}
            }
            r.set("device_check_have_data", json.dumps(device_check_have_data))

        tasks.publish_message(
            {
                "type": "kill",
                "id": serializer.data["id"],
                "ip_agent": serializer.data["ip"],
                "protocol": serializer.data["protocol"],
                "ip_django": os.getenv("LAN_IP"),
                "username": serializer.data["username"],
                "password": serializer.data["password"],
                "OS": serializer.data["os"] or "unknown",
                "port": serializer.data["port"],
                "agent_name": serializer.data["agent_name"],
                "hwagent": serializer.data["hwagent"],
                "device_type": serializer.data["device_type"],
            }
        )
        # kill agent thì cũng phải reset trạng thái trace syscall

        device.tracing_syscall = ""
        device.save()
        return Response({"status": "doing task..."})


class update_devices_by_macaddr(generics.GenericAPIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = DevicesSerializer
    queryset = Devices.objects.all()

    def get_object(self, mac_addr):
        try:
            return Devices.objects.get(mac_addr=mac_addr)
        except Devices.DoesNotExist:
            raise Http404

    # khong ghi ro method
    # @swagger_auto_schema(
    #     method='patch',
    #     request_body=UpdateDeviceSerializer,
    #     responses={201: openapi.Response(
    #         'DevicesSerializer', DevicesSerializer)}
    # )
    # @api_view(['PATCH'])
    def patch(self, request, mac_addr, format=None):
        xxx_data = json.loads(request.query_params["data"])
        devices = self.get_object(mac_addr)
        serializer = DevicesSerializer(devices, data=xxx_data)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class update_devices_by_ip(generics.GenericAPIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = UpdateDeviceSerializer
    queryset = Devices.objects.all()

    def get_object(self, ip):
        try:
            return Devices.objects.get(ip=ip)
        except Devices.DoesNotExist:
            raise Http404

    # khong ghi ro method
    # @swagger_auto_schema(
    #     method='patch',
    #     request_body=UpdateDeviceSerializer,
    #     responses={201: openapi.Response(
    #         'DevicesSerializer', DevicesSerializer)}
    # )
    # @api_view(['PATCH'])
    def patch(self, request, ip, format=None):
        xxx_data = json.loads(request.query_params["data"])
        # logger.error("error in update_devices...... request = %r "%(request))
        devices = self.get_object(ip)
        # devices
        serializer = UpdateDeviceSerializer(devices, data=xxx_data)
        # serializer = UpdateDeviceSerializer(devices)
        if serializer.is_valid():
            serializer.save()
            logger.info("done updated device by ip = %r" % (ip))
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class install_agent(APIView):
    permission_classes = [permissions.IsAdminUser]

    ip_agent = openapi.Parameter(
        "ip_agent", openapi.IN_QUERY, description="ip_agent", type=openapi.TYPE_STRING
    )
    ip_django = openapi.Parameter(
        "ip_django", openapi.IN_QUERY, description="ip_django", type=openapi.TYPE_STRING
    )
    ip_sock_serv = openapi.Parameter(
        "ip_sock_serv",
        openapi.IN_QUERY,
        description="ip_sock_serv",
        type=openapi.TYPE_STRING,
    )
    username = openapi.Parameter(
        "username", openapi.IN_QUERY, description="username", type=openapi.TYPE_STRING
    )
    password = openapi.Parameter(
        "password", openapi.IN_QUERY, description="password", type=openapi.TYPE_STRING
    )
    Devices_Serializer = openapi.Response("DevicesSerializer", DevicesSerializer)
    avatar = openapi.Parameter(
        "avatar", openapi.IN_QUERY, description="avatar", type=openapi.TYPE_FILE
    )

    os = openapi.Parameter(
        "os system", openapi.IN_QUERY, description="os system", type=openapi.TYPE_STRING
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[
            ip_agent,
            ip_django,
            ip_sock_serv,
            username,
            password,
            avatar
            #   , os
        ],
        responses={200: Devices_Serializer},
    )
    @api_view(["GET"])
    def get(self, request, format=None):
        tasks.publish_message(
            {
                "type": "install",
                "ipagent": request.query_params["ipagent"],
                "ip_django": request.query_params["ipdjango"],
                "ipsockserv": request.query_params["ipsockserv"],
                "username": request.query_params["username"],
                "password": request.query_params["password"],
                "os": request.query_params["os"],
                "port": request.query_params["port"],
            }
        )
        return Response(status=201, data="Ok")


class SyscallBatchView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = SyscallListSerializer
    queryset = SyscallList.objects.all()

    def post(self, request):
        xxx_data = json.loads(request.data)
        syscall_batch_as_list = xxx_data
        # SyscallList.objects.bulk_create(
        #     SyscallList(
        #         pid=each['pid'],
        #         syscall=each['syscall'],
        #         params=each['params'],
        #         device_id=Devices.objects.get(id=each['id']))
        #     for each in syscall_batch_as_list
        #     )
        for each_record in syscall_batch_as_list:
            syscall_object = SyscallList(
                pid=each_record["pid"],
                syscall=each_record["syscall"],
                params=each_record["params"],
                device_id=Devices.objects.get(id=each_record["id"]),
            )
            syscall_object.save()
            # goi ham nay hay bi do nen a cmt lai
            # save_syscall(sender=SyscallList,instance=syscall_object)
        return Response(status=201, data={"xxx_data"})


class BlackListIPView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = BlackListIPallSerializer
    queryset = BlackListIP.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend]
    # filterset_fields = ['current',
    #                     'page_size',
    #                     ]

    page = openapi.Parameter(
        "current",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    id = openapi.Parameter(
        "id",
        openapi.IN_QUERY,
        description="ID của thiết bị",
        type=openapi.TYPE_STRING,
        required=True,
    )
    ip = openapi.Parameter(
        "ip",
        openapi.IN_QUERY,
        description="IP",
        type=openapi.TYPE_STRING,
    )
    page_size = openapi.Parameter(
        "page_size",
        openapi.IN_QUERY,
        description="Số lượng mỗi trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    page = openapi.Parameter(
        "current",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    url = openapi.Parameter(
        "url",
        openapi.IN_QUERY,
        description="Địa chỉ trang web",
        type=openapi.TYPE_STRING,
        # required=True
    )

    def auto_update_black_list():
        try:
            req = requests.get(
                url=f'{os.getenv("CENTER_SERVER_IP")}/BlackListIP/',
                headers=get_headers(),
            )
            data = req.json()["results"]
            recent_id = len(BlackListIP.objects.all()) + 1
            for item in data:
                if (
                    BlackListIP.objects.all().filter(ip=item["ip"]).first() == None
                    and WhiteListIP.objects.all().filter(ip=item["ip"]).first() == None
                    and validateip.is_ipv4(item["ip"])
                ):
                    ip = BlackListIP(
                        id=(recent_id),
                        ip=f'{item["ip"]}',
                        url=f'{item["url"]}',
                        created=f'{item["created"]}',
                    )
                    ip.save()
                    recent_id += 1
            return Response({"msg": "success"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id, ip, page, page_size],
    )
    @action(detail=False, methods=["get"])
    def find_black_list_IP_in_device(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                id=request.query_params.get("id", None),
            ).first()

            listIP = data.blacklistip.filter(
                ip__icontains=request.query_params.get("ip", None)
            )
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(listIP, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
            return Response(json.loads(result))

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id, page, page_size],
    )
    @action(detail=False, methods=["get"])
    def get_black_list_IP_in_device(self, request):
        try:
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]

            device = Devices.objects.filter(id=request.query_params["id"]).first()
            # paginator = Paginator(device, page_size)
            data = device.blacklistip.all()
            if request.query_params.get("ip", None) != None:
                data = data.filter(ip=request.query_params["ip"])
            if request.query_params.get("url", None) != None:
                data = data.filter(url=request.query_params["url"])
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    # @swagger_auto_schema(
    #     method='get',
    #     manual_parameters=[id, page_size, page],

    # )
    # @action(detail=False, methods=['get'])
    # def get_black_list_IP_in_device(self, request):
    #     try:

    #         device = Devices.objects.filter(
    #                     id=request.query_params['id']).first()

    #         data = device.blacklistip.all()
    #         page_size = request.query_params['page_size']
    #         page = request.query_params['current']
    #         paginator = Paginator(data, page_size)
    #         serializer = self.get_serializer(
    #             paginator.get_page(page), many=True)
    #         return Response({
    #             'results': serializer.data,
    #             'count': paginator.count,
    #         })

    #     except Exception as e:
    #         logging.error(e)
    #         return Response(data={'status': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, page_size, page],
    )
    @action(detail=False, methods=["get"])
    def find_device_by_black_list_IP(self, request):
        try:
            data = Devices.objects.filter(
                blacklistip__ip__contains=request.query_params.get("ip", None),
            )
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[page, page_size],
    )
    @action(detail=False, methods=["get"])
    def get_all_black_list_IP_in_device(self, request):
        try:
            data = BlackListIP.objects.all()
            if request.query_params.get("ip", None) != None:
                data = data.filter(ip=request.query_params["ip"])
            if request.query_params.get("url", None) != None:
                data = data.filter(url=request.query_params["url"])
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            # if request.query_params
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    # @swagger_auto_schema(
    #     method='get',
    #     manual_parameters=[id, id, page_size, page],
    # )
    # @action(detail=False, methods=['get'])
    # def find_black_list_IP_in_device(self, request):
    #     try:
    #         # device = Devices.objects.get(pk=id)
    #         data = Devices.objects.filter(
    #                     id= request.query_params.get('id', None),
    #         ).first()
    #         listIP = data.blacklistip.filter(
    #             ip__icontains=request.query_params.get('ip', None)
    #             )
    #         page_size = request.query_params['page_size']
    #         page = request.query_params['current']
    #         paginator = Paginator(listIP, page_size)
    #         serializer = self.get_serializer(
    #             paginator.get_page(page), many=True)
    #         return Response({
    #             'results': serializer.data,
    #             'count': paginator.count,
    #         })
    #         # return Response({"message": "Done"})
    #     except Exception as e:
    #         logging.error(e)
    #         return Response(data={'status': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, url, page, page_size],
    )
    @action(detail=False, methods=["get"])
    def find_black_list_IP(self, request):
        try:
            data = BlackListIP.objects.all()
            if request.query_params.get("ip", None) != None:
                data = data.filter(ip__icontains=request.query_params.get("ip", None))
            if request.query_params.get("url", None) != None:
                data = data.filter(url__icontains=request.query_params.get("url", None))
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class WhiteListIPView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = WhiteListIPallSerializer
    queryset = WhiteListIP.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend]
    # # filterset_fields = ['current',
    # #                     'page_size',
    # #                     ]
    page = openapi.Parameter(
        "current",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    ip = openapi.Parameter(
        "ip",
        openapi.IN_QUERY,
        description="IP",
        type=openapi.TYPE_STRING,
        # required=True
    )
    id = openapi.Parameter(
        "id",
        openapi.IN_QUERY,
        description="ID của thiết bị",
        type=openapi.TYPE_STRING,
        required=True,
    )
    page_size = openapi.Parameter(
        "page_size",
        openapi.IN_QUERY,
        description="Số lượng mỗi trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    url = openapi.Parameter(
        "url",
        openapi.IN_QUERY,
        description="Địa chỉ trang web",
        type=openapi.TYPE_STRING,
        # required=True
    )

    def auto_update_white_list():
        try:
            req = requests.get(
                url=f'{os.getenv("CENTER_SERVER_IP")}/WhiteListIP/',
                headers=get_headers(),
            )
            data = req.json()["results"]
            recent_id = len(WhiteListIP.objects.all()) + 1
            for item in data:
                if (
                    BlackListIP.objects.all().filter(ip=item["ip"]).first() == None
                    and WhiteListIP.objects.all().filter(ip=item["ip"]).first() == None
                    and validateip.is_ipv4(item["ip"])
                ):
                    ip = WhiteListIP(
                        id=(recent_id),
                        ip=f'{item["ip"]}',
                        url=f'{item["url"]}',
                        created=f'{item["created"]}',
                    )
                    ip.save()
                    recent_id += 1
            return Response({"msg": "success"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id, page, page_size],
    )
    @action(detail=False, methods=["get"])
    def get_white_list_IP_in_device(self, request):
        try:
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]

            device = Devices.objects.filter(id=request.query_params["id"]).first()
            # paginator = Paginator(device, page_size)
            data = device.whitelistip.all()
            if request.query_params.get("ip", None) != None:
                data = data.filter(ip=request.query_params["ip"])
            if request.query_params.get("url", None) != None:
                data = data.filter(url=request.query_params["url"])
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[page, page_size],
    )
    @action(detail=False, methods=["get"])
    def get_all_white_list_IP_in_device(self, request):
        try:
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            data = WhiteListIP.objects.all()
            if request.query_params.get("ip", None) != None:
                data = data.filter(ip=request.query_params["ip"])
            if request.query_params.get("url", None) != None:
                data = data.filter(url=request.query_params["url"])
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, page_size, page],
    )
    @action(detail=False, methods=["get"])
    def find_device_by_white_list_IP(self, request):
        try:
            data = Devices.objects.filter(
                whitelistip__ip__contains=request.query_params.get("ip", None),
            )
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id, ip, page, page_size],
    )
    @action(detail=False, methods=["get"])
    def find_white_list_IP_in_device(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                id=request.query_params.get("id", None),
            ).first()
            listIP = data.whitelistip.filter(
                ip__icontains=request.query_params.get("ip", None)
            )
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(listIP, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, url, page, page_size],
    )
    @action(detail=False, methods=["get"])
    def find_white_list_IP(self, request):
        try:
            data = WhiteListIP.objects.all()
            if request.query_params.get("ip", None) != None:
                data = data.filter(ip__icontains=request.query_params.get("ip", None))
            if request.query_params.get("url", None) != None:
                data = data.filter(url__icontains=request.query_params.get("url", None))
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(data, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class DevicesViewSet(viewsets.ModelViewSet, generics.CreateAPIView):
    queryset = Devices.objects.all().order_by("created")
    # queryset = get_devices.delay().get()
    # pagination_class = None
    serializer_class = DevicesSerializer
    # permission_classes = [permissions.IsAuthenticated]
    parser_classes = [FormParser, MultiPartParser, JSONParser]

    # logger.error('================================', get_devices.delay())
    filter_backends = [DjangoFilterBackend]
    filterset_fields = [
        "username",
        "ip",
        "agentInstalled",
        "mac_addr",
        "tracing_syscall",
        "id",
        "protocol",
        "name",
        "device_type",
    ]

    def list(self, request, *args, **kwargs):
        devices = Devices.objects.all()
        # serializer = DevicesSerializer(devices, many=True)
        total = len(devices)
        # total_route = len(Devices.objects.filter(device_type='router'))
        # total_ip_cam = len(Devices.objects.filter(device_type='ip_cam'))
        # total_smartbox = len(Devices.objects.filter(device_type='smartbox'))
        # total_getway = len(Devices.objects.filter(device_type='gateway'))
        # total_hardware = len(Devices.objects.filter(device_type='hardware'))
        # try:
        #     host=os.environ.get("ELASTIC_HOST")

        #     es = Elasticsearch(hosts=host)
        #     doc = {
        #         "timestamp":datetime.now(timezone.utc),
        #         "total":total,
        #         "router": total_route,
        #         "ip_cam": total_ip_cam,
        #         "smartbox": total_smartbox,
        #         "gateway": total_getway,
        #         "hardware": total_hardware,
        #     }
        #     res = es.index(index="device_kc_info", document=doc)
        resData = []
        for device in devices:
            serializer = DevicesSerializer(device).data
            resData.append(serializer)
        return Response({"count": total, "results": resData}, status=status.HTTP_200_OK)
        # except Exception as ex:
        #     logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")
        # return Response(serializer.data)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance)
        data = serializer.data
        if r.get(f"cpu_ram_deivce_{data['id']}"):
            data_device_redis = json.loads(
                r.get(f"cpu_ram_deivce_{data['id']}").decode()
            )
            resource_info = data_device_redis["resource_info"]
            agentInstalled = data_device_redis["agentInstalled"]
            data["resource_info"] = resource_info
            data["agentInstalled"] = agentInstalled
        return Response(data)

    # def update(self, request, *args, **kwargs):
    #     partial = kwargs.pop('partial', False)
    #     instance = self.get_object()
    #     serializer = self.get_serializer(instance, data=request.data, partial=partial)
    #     serializer.is_valid(raise_exception=True)
    #     self.perform_update(serializer)
    #     data = serializer.data
    #     resource_info =  {
    #                         "SPEED_UP": random.randint(3, 9),
    #                         "SPEED_DOWN": random.randint(3, 9),
    #                         "timestamps": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
    #                         "PERCENT_CPU": random.randint(3, 9),
    #                         "PERCENT_RAM": random.randint(3, 9),
    #                         "OPENING_PORT": []
    #                     }
    #     agentInstalled = True
    #     tracing_network = True
    #     data["resource_info"]=resource_info
    #     data["agentInstalled"]=agentInstalled
    #     data["tracing_network"]=tracing_network
    #     return Response(data)

    @action(detail=False, methods=["post"])
    def patch_many_device(self, request):
        return Response(status=status.HTTP_204_NO_CONTENT)

    @swagger_auto_schema(
        method="get",
        # responses=openapi.Response('DevicesSerializer', DevicesSerializer)
    )
    @action(detail=False, methods=["get"])
    def get_info_device_to_elastic(self, request, *args, **kwargs):
        self.send_elastic_info_all_device()
        return Response(status=status.HTTP_200_OK)

    def send_elastic_info_all_device(self):
        total = len(Devices.objects.all())
        total_route = len(Devices.objects.filter(device_type="router"))
        total_ip_cam = len(Devices.objects.filter(device_type="ip_cam"))
        total_smartbox = len(Devices.objects.filter(device_type="smartbox"))
        total_getway = len(Devices.objects.filter(device_type="gateway"))
        total_hardware = len(Devices.objects.filter(device_type="hardware"))
        # elastic local
        try:
            host = os.environ.get("ELASTIC_HOST")
            es = Elasticsearch(hosts=host)
            doc = {
                "timestamp": datetime.now(timezone.utc),
                "total": total,
                "router": total_route,
                "ip_cam": total_ip_cam,
                "smartbox": total_smartbox,
                "gateway": total_getway,
                "hardware": total_hardware,
                "local": os.environ.get("LOCAL_SITE"),
                "type_log": "device_kc_info",
            }
            res = es.index(index="demo-kc", document=doc)
        except Exception as ex:
            logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")
        # elasttic center
        try:
            host = os.environ.get("CENTER_SERVER_ELASTIC")
            es = Elasticsearch(hosts=host)
            doc = {
                "timestamp": datetime.now(timezone.utc),
                "total": total,
                "router": total_route,
                "ip_cam": total_ip_cam,
                "smartbox": total_smartbox,
                "gateway": total_getway,
                "hardware": total_hardware,
                "local": os.environ.get("LOCAL_SITE"),
                "type_log": "device_kc_info",
            }
            res = es.index(index="demo-kc", document=doc)
        except Exception as ex:
            logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")

    def create(self, request, *args, **kwargs):
        serializer = DevicesSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            data_device_norm = {"device": serializer.data["id"]}
            serializer_device_norm = DeviceNormSerializer(data=data_device_norm)
            if serializer_device_norm.is_valid():
                serializer_device_norm.save()
                action_device_norm("create", serializer_device_norm.data)
                action_check_device_norm(serializer_device_norm.data)
            self.send_elastic_info_all_device()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def destroy(self, request, pk, format=None):
        devices = get_object_or_404(Devices.objects.all(), pk=pk)
        devices.delete()
        self.send_elastic_info_all_device()
        return Response(status=status.HTTP_204_NO_CONTENT)

    def update_rulebase_hw(self, pk):
        current_device = Devices.objects.filter(id=pk).first()
        if current_device.hwagent:
            tasks.publish_message(
                {
                    "type": "update rulebase hw",
                    "id": current_device.id,
                    "ip_agent": current_device.ip,
                    "protocol": current_device.protocol,
                    "ip_django": os.getenv("LAN_IP"),
                    "ip_sock_serv": os.getenv("LAN_SOCKET_SERVER_IP")
                    if not current_device.from_internet
                    else os.getenv("PUBLIC_SOCKET_SERVER"),
                    "ip_tftp_serv": os.getenv("LAN_SOCKET_SERVER_IP")
                    if not current_device.from_internet
                    else os.getenv("PUBLIC_TFTP_SERVER"),
                    "username": current_device.username,
                    "password": current_device.password,
                    "OS": current_device.os,
                    "port": current_device.port,
                    "remote_port": current_device.remote_port,
                    "hwagent": current_device.hwagent,
                }
            )
        return "doing task update rulebase hardware agent"

    id = openapi.Parameter(
        "id",
        openapi.IN_QUERY,
        description="ID Thiết bị",
        required=True,
        type=openapi.TYPE_STRING,
    )
    ip = openapi.Parameter(
        "ip",
        openapi.IN_QUERY,
        description="IP Thiết bị",
        required=True,
        type=openapi.TYPE_STRING,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id],
        # responses=openapi.Response('DevicesSerializer', DevicesSerializer)
    )
    @action(detail=False, methods=["get"])
    def reboot(self, request):
        try:
            current_device = Devices.objects.filter(
                id=request.query_params["id"]
            ).first()
            # Elastic local
            try:
                host = os.environ.get("ELASTIC_HOST")
                es = Elasticsearch(hosts=host)
                doc = {
                    "timestamp": datetime.now(timezone.utc),
                    "ip": current_device.ip,
                    "device_id": request.query_params["id"],
                    "type": "Khởi động lại",
                    "type_log": "alert_prevent_kc",
                }
                res = es.index(index="demo-kc", document=doc)
            except Exception as ex:
                logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")

            # Elastic center
            try:
                host = os.environ.get("CENTER_SERVER_ELASTIC")
                es = Elasticsearch(hosts=host)
                doc = {
                    "timestamp": datetime.now(timezone.utc),
                    "ip": current_device.ip,
                    "device_id": request.query_params["id"],
                    "type": "Khởi động lại",
                    "local": os.environ.get("LOCAL_SITE"),
                    "type_log": "alert_prevent_kc",
                }
                res = es.index(index="demo-kc", document=doc)
            except Exception as ex:
                logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")

            tasks.publish_message(
                {
                    "type": "reboot",
                    "id": current_device.id,
                    "ip_agent": current_device.ip,
                    "protocol": current_device.protocol,
                    "ip_django": os.getenv("LAN_IP"),
                    "username": current_device.username,
                    "password": current_device.password,
                }
            )
            return Response({"status": 200})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=["get"])
    def isolate(self, request):
        try:
            current_device = Devices.objects.filter(
                id=request.query_params["id"]
            ).first()
            # Elastic local
            try:
                host = os.environ.get("ELASTIC_HOST")
                es = Elasticsearch(hosts=host)
                doc = {
                    "timestamp": datetime.now(timezone.utc),
                    "ip": current_device.ip,
                    "device_id": request.query_params["id"],
                    "type": "Cách ly mạng",
                    "type_log": "alert_prevent_kc",
                }
                res = es.index(index="demo-kc", document=doc)
            except Exception as ex:
                logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")

            # Elastic center
            try:
                host = os.environ.get("CENTER_SERVER_ELASTIC")
                es = Elasticsearch(hosts=host)
                doc = {
                    "timestamp": datetime.now(timezone.utc),
                    "ip": current_device.ip,
                    "device_id": request.query_params["id"],
                    "type": "Cách ly mạng",
                    "local": os.environ.get("LOCAL_SITE"),
                    "type_log": "alert_prevent_kc",
                }
                res = es.index(index="demo-kc", document=doc)
            except Exception as ex:
                logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")

            tasks.publish_message(
                {
                    "type": "isolate",
                    "id": current_device.id,
                    "ip_agent": current_device.ip,
                    "protocol": current_device.protocol,
                    "ip_django": os.getenv("LAN_IP"),
                    "username": current_device.username,
                    "password": current_device.password,
                }
            )
            return Response({"status": 200})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=["post"])
    def post(self, request, *args, **kwargs):
        xxx_data = json.loads(request.query_params["data"])
        file_serializer = DevicesSerializer(data=xxx_data)
        if file_serializer.is_valid():
            file_serializer.save()
            return Response(file_serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response(file_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    # @swagger_auto_schema(
    #     method='patch',
    #     request_body=UpdateIPFrequencySerializer,
    #     # responses=openapi.Response('DevicesSerializer', DevicesSerializer)
    # )
    # # @api_view(['PATCH'])
    # @action(detail=False, methods=['patch'])
    # def add_dst_ip(self, request):
    #     try:
    #         current_device = Devices.objects.filter(
    #             ip=xxx_data['ip']).first()
    #         if (current_device.ips == []):
    #             current_device.ips = {}
    #         for ip in xxx_data['list_ip']:
    #             # Check trường hợp mặc định là [] -> đổi sang dict
    #             #  update tần suất
    #             if not (ip in current_device.ips):
    #                 current_device.ips[ip] = 1
    #             else:
    #                 current_device.ips[ip] = current_device.ips[ip] + 1
    #         current_device.save()
    #         return Response(DevicesSerializer(current_device).data)
    #     except Exception as e:
    #         logging.error(e)
    #         return Response(data={'status': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=DevicesSerializer,
        # responses=openapi.Response('DevicesSerializer', DevicesSerializer),
    )
    @action(detail=False, methods=["patch"])
    def update_device_by_ip(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            current_device = Devices.objects.filter(ip=xxx_data["ip"]).first()
            for key in xxx_data:
                setattr(current_device, key, xxx_data[key])

            current_device.save()
            return Response(DevicesSerializer(current_device).data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=DevicesSerializer,
        # responses=openapi.Response('DevicesSerializer', DevicesSerializer),
    )
    @action(detail=False, methods=["patch"])
    def update_resource_info(self, request):
        # queryset = Devices.objects.all().order_by('created')
        # for device in queryset.iterator():
        xxx_data = json.loads(request.query_params["data"])
        try:
            current_device = Devices.objects.filter(ip=xxx_data["ip"]).first()
            current_device.resource_info = xxx_data["resource_info"]
            current_device.save()
            return Response(DevicesSerializer(current_device).data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def auto_update_rule_snort(self):
        logging.error("!!!!!!!!!!ALO Update rule nao :3")
        tasks.publish_message({"type": "update_rule"})
        return Response({"status": "doing task..."})

    # @action(detail=False, methods=['patch'])
    def auto_update_password(self):
        queryset = Devices.objects.all().filter(autoUpdatePasswd=True)
        characters = string.ascii_letters + string.digits
        for device in queryset.iterator():
            autogen_password = "".join(random.choice(characters) for i in range(15))
            auto_gen = "clgt123456"
            serializer = DevicesSerializer(device)
            telnet_root_dir = "./kc-static-files/build"
            if not os.path.isdir(telnet_root_dir):
                os.makedirs(telnet_root_dir)
            createFileObject = CreateInfoFile(autogen_password, telnet_root_dir)
            createFileObject.do_create_file()

            tasks.publish_message(
                {
                    "type": "chpasswd",
                    "id": serializer.data["id"],
                    "ip_agent": serializer.data["ip"],
                    "ip_django": os.getenv("LAN_IP"),
                    "ip_sock_serv": os.getenv("LAN_SOCKET_SERVER_IP"),
                    "username": serializer.data["username"],
                    "password": serializer.data["password"],
                    "protocol": serializer.data["protocol"],
                    "new_password": autogen_password,
                }
            )
        return Response({"msg": "success"})

    @swagger_auto_schema(method="get")
    @action(detail=False, methods=["get"])
    def reset_status_device(self, request):
        try:
            queryset = Devices.objects.all()
            for device in queryset.iterator():
                device.agentInstalled = False
                device.tracing_network = False
                device.save()
            return Response(True)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(method="patch", request_body=InstallAgentSerializer)
    @action(detail=False, methods=["patch"])
    def install_list_agent(self, request):
        try:
            data = json.loads(request.body.decode("utf-8"))
            for dt in data["id"]:
                currentDevice = (
                    Devices.objects.filter(id=dt).order_by("-created").first()
                )
                tasks.publish_message(
                    {
                        "type": "install",
                        "id": currentDevice.id,
                        "ip_agent": currentDevice.ip,
                        "protocol": currentDevice.protocol,
                        "ip_django": os.getenv("LAN_IP"),
                        "ip_sock_serv": os.getenv("LAN_SOCKET_SERVER_IP")
                        if not currentDevice.from_internet
                        else os.getenv("PUBLIC_SOCKET_SERVER"),
                        "ip_tftp_serv": os.getenv("LAN_SOCKET_SERVER_IP")
                        if not currentDevice.from_internet
                        else os.getenv("PUBLIC_TFTP_SERVER"),
                        "username": currentDevice.username,
                        "password": currentDevice.password,
                        "OS": currentDevice.os,
                        "port": currentDevice.port,
                        "remote_port": currentDevice.remote_port,
                        "hwagent": currentDevice.hwagent,
                    }
                )
            return Response({"message": "handling.."})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(method="patch", request_body=InstallAgentSerializer)
    @action(detail=False, methods=["patch"])
    def kill_list_agent(self, request):
        try:
            data = json.loads(request.body.decode("utf-8"))
            for dt in data["id"]:
                currentDevice = (
                    Devices.objects.filter(id=dt).order_by("-created").first()
                )
                tasks.publish_message(
                    {
                        "type": "kill",
                        "id": currentDevice.id,
                        "ip_agent": currentDevice.ip,
                        "protocol": currentDevice.protocol,
                        "ip_django": os.getenv("LAN_IP"),
                        "username": currentDevice.username,
                        "password": currentDevice.password,
                        "OS": currentDevice.os or "unknown",
                        "port": currentDevice.port,
                        "hwagent": currentDevice.hwagent,
                    }
                )
            return Response({"message": "handling.."})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=BlackListIPSerializer,
    )
    @action(detail=False, methods=["patch"])
    def add_black_list_IP(self, request):
        try:
            xxx_data = request.data
            new = []
            white = []
            normal = []
            data = xxx_data["blacklistip"]
            messss = "doing task update rulebase hardware agent"
            for id in xxx_data["id"]:
                for dt in data:
                    device = Devices.objects.filter(id=id).first()
                    whitelist = device.whitelistip.all()
                    stt = True
                    ip = BlackListIP.objects.filter(id=dt).first()
                    for dt_white in whitelist:
                        if ip.ip == dt_white.ip:
                            white.append(dt_white.ip)
                            if stt:
                                messss = (
                                    dt_white.ip
                                    + " đã có trong WhiteList không thể thêm vào BlackList"
                                )
                                stt = False
                            else:
                                messss = dt_white.ip + " " + messss
                    if stt:
                        ip = BlackListIP.objects.filter(id=dt).first()
                        normal.append(dt)
                        device.blacklistip.add(ip)
                        device.save()

                        # Elastic local
                        try:
                            host = os.environ.get("ELASTIC_HOST")
                            es = Elasticsearch(hosts=host)
                            doc = {
                                "timestamp": datetime.now(timezone.utc),
                                "ip": device.ip,
                                "device_id": id,
                                "type": "Thêm IP/URL vào blacklist",
                                "type_log": "alert_prevent_kc",
                            }
                            res = es.index(index="demo-kc", document=doc)
                        except Exception as ex:
                            logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")

                        # Elastic center
                        try:
                            host = os.environ.get("CENTER_SERVER_ELASTIC")
                            es = Elasticsearch(hosts=host)
                            doc = {
                                "timestamp": datetime.now(timezone.utc),
                                "ip": device.ip,
                                "device_id": id,
                                "type": "Thêm IP/URL vào blacklist",
                                "local": os.environ.get("LOCAL_SITE"),
                                "type_log": "alert_prevent_kc",
                            }
                            res = es.index(index="demo-kc", document=doc)
                        except Exception as ex:
                            logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")
                rep = self.update_rulebase_hw(id)
            return Response({"message": messss})

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=WhiteListIPSerializer,
    )
    @action(detail=False, methods=["patch"])
    def add_white_list_IP(self, request):
        try:
            xxx_data = request.data
            new = []
            black = []
            data = xxx_data["whitelistip"]
            messss = "Done"
            for id in xxx_data["id"]:
                for dt in data:
                    device = Devices.objects.filter(id=id).first()
                    blacklist = device.blacklistip.all()
                    stt = True
                    ip = WhiteListIP.objects.filter(id=dt).first()
                    for dt_black in blacklist:
                        if ip.ip == dt_black.ip:
                            black.append(dt_black.id)
                            ip_black = BlackListIP.objects.filter(
                                id=dt_black.id
                            ).first()
                            device.blacklistip.remove(ip_black)

                            if stt:
                                messss = (
                                    dt_black.ip + " có trong BlackList và đã bị xóa"
                                )
                                stt = False
                            else:
                                messss = dt_black.ip + " " + messss

                    device.whitelistip.add(ip)
                    device.save()
                rep = self.update_rulebase_hw(id)
            return Response({"message": messss})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=WhiteListIPSerializer,
    )
    @action(detail=False, methods=["patch"])
    def remove_white_list_IP(self, request):
        try:
            xxx_data = request.data
            data = xxx_data["whitelistip"]
            for id in xxx_data["id"]:
                device = Devices.objects.filter(id=id).first()
                for dt in data:
                    ip = WhiteListIP.objects.filter(id=dt).first()
                    device.whitelistip.remove(ip)
                    device.save()
            return Response({"message": "handling.."})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=BlackListIPSerializer,
    )
    @action(detail=False, methods=["patch"])
    def remove_black_list_IP(self, request):
        try:
            xxx_data = request.data
            data = xxx_data["blacklistip"]
            for id in xxx_data["id"]:
                for dt in data:
                    device = Devices.objects.filter(id=id).first()
                    ip = BlackListIP.objects.filter(id=dt).first()
                    device.blacklistip.remove(ip)
                    device.save()
                data = self.update_rulebase_hw(id)
            return Response({"message": "Done"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    url = openapi.Parameter(
        "url", openapi.IN_QUERY, description="url", type=openapi.TYPE_STRING
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, url],
    )
    @action(detail=False, methods=["get"])
    def find_device_by_black_list_IP(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                blacklistip__ip__contains=request.query_params.get("ip", None),
            )
            result = serializers.serialize("json", data)
            return Response(json.loads(result))
            # return Response({"message": "Done"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, url],
    )
    @action(detail=False, methods=["get"])
    def find_device_white_list_IP(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                whitelistip__ip__contains=request.query_params.get("ip", None),
            )
            result = serializers.serialize("json", data)
            return Response(json.loads(result))
            # return Response({"message": "Done"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id, ip, url],
    )
    @action(detail=False, methods=["get"])
    def find_white_list_IP_in_device(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                id=request.query_params.get("id", None),
            ).first()
            listIP = data.whitelistip.filter(
                ip__icontains=request.query_params.get("ip", None)
            )[:20]
            result = serializers.serialize("json", listIP)
            return Response(json.loads(result))
            # return Response({"message": "Done"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    page = openapi.Parameter(
        "current",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    page_size = openapi.Parameter(
        "page_size",
        openapi.IN_QUERY,
        description="Số lượng mỗi trang",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id, ip, url],
    )
    @action(detail=False, methods=["get"])
    def find_black_list_IP_in_device(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                id=request.query_params.get("id", None),
            ).first()
            listIP = data.blacklistip.filter(
                ip__icontains=request.query_params.get("ip", None)
            )[:20]
            result = serializers.serialize("json", listIP)
            return Response(json.loads(result))

        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[id],
    )
    @action(detail=False, methods=["get"])
    def find_black_list_IP_in_device_to_Ip(self, request):
        try:
            # device = Devices.objects.get(pk=id)
            data = Devices.objects.filter(
                id=request.query_params.get("id", None),
            ).first()
            listIP = data.blacklistip.all()
            rel = []
            for dt in listIP:
                rel.append(dt.ip)
            return Response({"data": rel})
            # return Response({"message": "Done"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class ChangePwdView(generics.GenericAPIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = UpdateDeviceTelnetPasswordSerializer
    queryset = Devices.objects.all()

    def get_object(self, pk):
        try:
            return Devices.objects.get(pk=pk)
        except Devices.DoesNotExist:
            raise Http404

    def patch(self, request, pk, format=None):
        xxx_data = request.data
        device = self.get_object(pk)
        serializer = DevicesSerializer(device, context={"request": request})
        telnet_root_dir = "./kc-static-files/build"
        if not os.path.isdir(telnet_root_dir):
            os.makedirs(telnet_root_dir)

        createFileObject = CreateInfoFile(xxx_data["password"], telnet_root_dir)
        createFileObject.do_create_file()

        tasks.publish_message(
            {
                "type": "chpasswd",
                "id": serializer.data["id"],
                "ip_agent": serializer.data["ip"],
                "ip_django": os.getenv("LAN_IP"),
                "ip_sock_serv": os.getenv("LAN_SOCKET_SERVER_IP"),
                "username": serializer.data["username"],
                "password": serializer.data["password"],
                "protocol": serializer.data["protocol"],
                "new_password": xxx_data["password"],
            }
        )
        # device.password = xxx_data['password']
        # device.save()
        return Response({"status": "doing task..."})

    def put(self, request, pk, format=None):
        xxx_data = request.data
        device = self.get_object(pk)
        serializer = DevicesSerializer(device, context={"request": request})
        device.password = xxx_data["password"]
        device.save()
        return Response({"status": "done chpasswd..."})


class StandardResultsSetPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = "page_size"
    max_page_size = 1000


class AlertsViewSet(viewsets.ModelViewSet):
    queryset = Alerts.objects.all().order_by("-timestamp", "status")
    # queryset = get_devices.delay().get()
    serializer_class = AlertsSerializer
    permission_classes = [permissions.IsAuthenticated | HasAPIKey]
    # logger.error('================================', get_devices.delay())
    pagination_class = StandardResultsSetPagination

    filter_backends = [DjangoFilterBackend]
    filterset_fields = [
        "ip",
        "timestamp",
        "hash",
        "pid",
        "message",
        "type",
        "status",
        "device",
    ]

    ip = openapi.Parameter(
        "ip",
        openapi.IN_QUERY,
        description="IP thiết bị",
        type=openapi.TYPE_STRING,
    )
    device = openapi.Parameter(
        "device",
        openapi.IN_QUERY,
        description="ID thiết bị",
        type=openapi.TYPE_STRING,
    )
    page = openapi.Parameter(
        "current",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    page_size = openapi.Parameter(
        "page_size",
        openapi.IN_QUERY,
        description="Số lượng mỗi trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    Type = openapi.Parameter(
        "type",
        openapi.IN_QUERY,
        description="Loại thông báo",
        type=openapi.TYPE_STRING,
        # required = True
    )
    status = openapi.Parameter(
        "status",
        openapi.IN_QUERY,
        description="Trạng thái",
        type=openapi.TYPE_STRING,
        # required = True
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[ip, page, page_size, Type, status],
        # responses=openapi.Response('AlertsSerializer', AlertsSerializer),
    )
    @action(detail=False, methods=["get"])
    def GetAlertsByIP(self, request):
        try:
            type = request.query_params.get("type", None)
            statusValue = request.query_params.get("status", None)
            filters = {"ip": request.query_params["ip"]}
            if type:
                filters["type"] = type
            if statusValue:
                filters["status"] = statusValue
            result = Alerts.objects.filter(**filters).order_by("-timestamp")

            page_size = request.query_params["page_size"]
            page = request.query_params["current"]

            paginator = Paginator(result, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="device_id",
        required=True,
        type=openapi.TYPE_STRING,
        # required = True
    )
    pid = openapi.Parameter(
        "pid",
        openapi.IN_QUERY,
        description="pid",
        required=True,
        type=openapi.TYPE_STRING,
        # required = True
    )
    alert_id = openapi.Parameter(
        "alert_id",
        openapi.IN_QUERY,
        description="alert_id",
        required=True,
        type=openapi.TYPE_STRING,
        # required = True
    )
    message = openapi.Parameter(
        "message",
        openapi.IN_QUERY,
        description="message",
        type=openapi.TYPE_STRING,
        # required = True
    )
    all_device = openapi.Parameter(
        "all_device",
        openapi.IN_QUERY,
        description="all_device",
        type=openapi.TYPE_BOOLEAN,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        # responses=openapi.Response('AlertsSerializer', AlertsSerializer),
        manual_parameters=[device_id, pid, alert_id],
    )
    @action(detail=False, methods=["get"])
    def kill_process(self, request):
        try:
            currentDevice = (
                Devices.objects.filter(id=request.query_params["device_id"])
                .order_by("-created")
                .first()
            )

            tasks.publish_message(
                {
                    "type": "kill_process",
                    "id": currentDevice.id,
                    "ip_agent": currentDevice.ip,
                    "ip_django": os.getenv("LAN_IP"),
                    "pid": request.query_params["pid"],
                    "alert_id": request.query_params["alert_id"],
                    "username": currentDevice.username,
                    "password": currentDevice.password,
                    "protocol": currentDevice.protocol,
                }
            )
            return Response({"message": "handling.."})

        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        manual_parameters=[message, ip, page_size, page, all_device],
        request_body=AlertsExportSerializer,
    )
    @action(detail=False, methods=["patch"])
    def search_alert(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            result = Alerts.objects.all()
            if request.query_params.get("ip", None) != None:
                result = result.filter(ip__icontains=request.query_params["ip"])
            if request.query_params.get("message", None) != None:
                result = result.filter(
                    message__icontains=request.query_params["message"]
                )
            if request.query_params.get("address", None) != None:
                result = result.filter(
                    address__icontains=request.query_params["address"]
                )

            if xxx_data.get("timestamp", None) != None:
                # aaa = datetime.strptime(xxx_data.get('timestamp', None)[0], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                # bbb = datetime.strptime(xxx_data.get('timestamp', None)[1], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                result = result.filter(
                    timestamp__range=[
                        xxx_data.get("timestamp", None)[0],
                        xxx_data.get("timestamp", None)[1],
                    ]
                )
            i = 0
            xyz = result
            if request.query_params["all_device"] == "false":
                if xxx_data.get("device", None) != None:
                    for id in xxx_data.get("device", None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0

            if xxx_data.get("type", None) != None:
                for id in xxx_data.get("type", None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = abc | result.filter(type=id)
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get("status", None) != None:
                for id in xxx_data.get("status", None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz

            page_size = request.query_params["page_size"]
            page = request.query_params["current"]
            paginator = Paginator(xyz, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )

        except Exception as e:
            logging.error(e)
            return Response(
                data={"status": False, "Log bug": str(e)},
                status=status.HTTP_400_BAD_REQUEST,
            )

    @swagger_auto_schema(
        method="patch",
        request_body=AlertsExportSerializer,
        manual_parameters=[message, ip, all_device],
        # manual_parameters=[message],
    )
    @action(detail=False, methods=["patch"])
    def exportXLSX(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            sheet1 = workbook.add_worksheet("Sheet1")
            # sheet1 = workbook.add_sheet('Sheet1', cell_overwrite_ok = True)
            col_stt = 1
            rol = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N"]
            rol_stt = 0
            for i in xxx_data["status"]:
                sheet1.write(
                    # rol[rol_stt]
                    "A" + str(col_stt),
                    "Số lần thực hiện của thể loại : " + i,
                )
                # col_stt = col_stt + 1
                sheet1.write(
                    # rol[rol_stt]
                    "B" + str(col_stt),
                    "Số lần thực hiện là "
                    + str(Alerts.objects.filter(status=i).count()),
                )
                col_stt = col_stt + 1
            for y in xxx_data["type"]:
                sheet1.write(
                    # rol[rol_stt]
                    "A" + str(col_stt),
                    "Số lần thực hiện của trạng thái: " + y,
                )
                # col_stt = col_stt + 1
                sheet1.write(
                    rol[rol_stt + 1] + str(col_stt),
                    "Số lần thực hiện là " + str(Alerts.objects.filter(type=y).count()),
                )
                col_stt = col_stt + 1
            sheet2 = workbook.add_worksheet("Sheet2")
            result = Alerts.objects.all()

            if request.query_params.get("ip", None) != None:
                result = result.filter(ip__icontains=request.query_params["ip"])

            if request.query_params.get("message", None) != None:
                result = result.filter(
                    message__icontains=request.query_params["message"]
                )
            if request.query_params.get("address", None) != None:
                result = result.filter(
                    address__icontains=request.query_params["address"]
                )

            if xxx_data.get("timestamp", None) != None:
                result = result.filter(
                    timestamp__gte=xxx_data.get("timestamp", None)[0],
                    timestamp__lte=xxx_data.get("timestamp", None)[1],
                )
            i = 0
            result_id = result
            if request.query_params["all_device"] == True:
                if xxx_data.get("id", None) != None:
                    for id in xxx_data.get("id", None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0
            if xxx_data.get("type", None) != None:
                for id in xxx_data.get("type", None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(type=id)).distinct()
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get("status", None) != None:
                for id in xxx_data.get("status", None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz

            header = [
                "STT",
                "ID",
                "IP",
                "Timestamp",
                "Hash",
                "PID",
                "Message",
                "Type",
                "Rule_base",
                "Log",
                "Status",
                "Device",
            ]
            col_stt = 1
            for rol_stt in range(12):
                sheet2.write(rol[rol_stt] + str(col_stt), header[rol_stt])
            col_stt = col_stt + 1
            for dt_res in result_id:
                sheet2.write(rol[0] + str(col_stt), col_stt - 1)
                sheet2.write(rol[1] + str(col_stt), dt_res.id)
                sheet2.write(rol[2] + str(col_stt), dt_res.ip)
                sheet2.write(
                    rol[3] + str(col_stt),
                    dt_res.timestamp.strftime("%Y-%m-%d %H:%M:%S"),
                )
                sheet2.write(rol[4] + str(col_stt), dt_res.hash)
                sheet2.write(rol[5] + str(col_stt), dt_res.pid)
                sheet2.write(rol[6] + str(col_stt), dt_res.message)
                sheet2.write(rol[7] + str(col_stt), dt_res.type)
                sheet2.write(rol[8] + str(col_stt), dt_res.rule_base)
                sheet2.write(rol[9] + str(col_stt), dt_res.log)
                sheet2.write(rol[10] + str(col_stt), dt_res.status)
                sheet2.write(rol[11] + str(col_stt), dt_res.device.id)
                col_stt = col_stt + 1
            workbook.close()

            # create a response
            response = HttpResponse(content_type="application/vnd.ms-excel")

            # tell the browser what the file is named
            response[
                "Content-Disposition"
            ] = 'attachment;filename="some_file_name.xlsx"'

            # put the spreadsheet data into the response
            response.write(output.getvalue())

            return response

        except Exception as e:
            logging.error(e)
            return Response(
                data={"status": False, "Log bug": e}, status=status.HTTP_400_BAD_REQUEST
            )


class ProcessHashViewSet(viewsets.ModelViewSet):
    queryset = ProcessHash.objects.all()
    serializer_class = ProcessHashSerializer
    permission_classes = [permissions.IsAuthenticated]

    sha1 = openapi.Parameter(
        "sha1",
        openapi.IN_QUERY,
        description="sha1",
        type=openapi.TYPE_STRING,
        required=True,
    )

    def create(self, request, *args, **kwargs):
        data = request.data
        # Elastic server
        try:
            host = os.environ.get("CENTER_SERVER_ELASTIC")
            es = Elasticsearch(hosts=host)
            doc = {
                "sha1": data.get("sha1"),
                "loaiMaDoc": data.get("loaiMaDoc"),
                "acc": data.get("acc"),
                "local": os.environ.get("LOCAL_SITE"),
                "type_log": "malware",
            }
            res = es.index(index="demo-kc", document=doc)
        except Exception as ex:
            logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")
        # Elastic local
        try:
            host = os.environ.get("ELASTIC_HOST")
            es = Elasticsearch(hosts=host)
            doc = {
                "sha1": data.get("sha1"),
                "loaiMaDoc": data.get("loaiMaDoc"),
                "acc": data.get("acc"),
                "type_log": "malware",
            }
            res = es.index(index="demo-kc", document=doc)
            return Response(data, status=status.HTTP_201_CREATED)
        except Exception as ex:
            logging.error("!!!!!!!!!!! Gửi lên ELK lỗi")
            return Response(
                "!!!!!!!!!!! Gửi lên ELK lỗi", status=status.HTTP_400_BAD_REQUEST
            )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[sha1],
        # responses=openapi.Response(
        #     'ProcessHashSerializer', ProcessHashSerializer),
        operation_description="get",
    )
    @action(detail=False, methods=["get"])
    def check_exist(self, request):
        try:
            result = ProcessHash.objects.get(sha1=request.query_params["sha1"])
            serializer = self.get_serializer(result)
            return Response(serializer.data)
        except ProcessHash.DoesNotExist:
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    # @api_view(['POST'])
    @action(detail=False, methods=["post"])
    def check_batch_exist(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            batch_of_pid_sha1 = xxx_data
            # query sha1_DB by sha1_agent
            batch_of_sha1 = [each["sha1"] for each in batch_of_pid_sha1]
            query_result_as_objects = ProcessHash.objects.filter(sha1__in=batch_of_sha1)
            sha1_acc_result_as_list = [
                {
                    "sha1": self.get_serializer(e).data["sha1"],
                    "acc": self.get_serializer(e).data["acc"],
                }
                for e in query_result_as_objects
            ]
            #   '???????', (batch_of_pid_sha1))
            result_pid_sha1 = []
            for each_pid_sha1 in batch_of_pid_sha1:
                exist_status = False
                pid = each_pid_sha1["pid"]
                sha1_agent = each_pid_sha1["sha1"]

                for each_sha1_acc in sha1_acc_result_as_list:
                    sha1_db = each_sha1_acc["sha1"]
                    acc = each_sha1_acc["acc"]
                    if sha1_agent == sha1_db:
                        result_pid_sha1.append(
                            {
                                "exist": 1,
                                "pid": pid,
                                "sha1": sha1_db,
                                "acc": acc,
                            }
                        )
                        exist_status = True
                        break  # luc nay loi o cho nay, vcl~
                # Nếu ko tồn tại thì trả về 0: Tức ko tồn tại
                if not exist_status:
                    result_pid_sha1.append(
                        {
                            "exist": 0,
                            "pid": pid,
                            "sha1": sha1_agent,
                            "acc": -1,
                        }
                    )
            return Response(data=result_pid_sha1, status=201)
        except Exception as e:
            logger.error(e)
            return Response(status=500)

    @action(detail=False, methods=["get"])
    def GetProcessListByDeviceId(self, request):
        try:
            init_data_malware()
            return Response(status=200)
        except Exception as e:
            return Response(status=500)


def add_process_list_x(data_new):
    f1 = open("./init_processlist.json")
    datas = json.load(f1)
    ok = False
    for data in datas:
        if int(data.get("device_id", 0)) == int(data_new.get("device_id")):
            data["process_list"] = data_new.get("process_list")
            ok = True
    if ok == False:
        datas.append(data_new)
    with open("./init_processlist.json", "w") as outfile:
        json.dump(datas, outfile)
    # pass


class ProcessListViewSet(viewsets.ModelViewSet):
    queryset = ProcessList.objects.all()
    serializer_class = ProcessListSerializer
    permission_classes = [permissions.IsAuthenticated]

    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="device_id",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @action(detail=False, methods=["post"])
    def push_many_process(self, request):
        return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
        # responses=openapi.Response(
        #     'ProcessListSerializer', ProcessListSerializer),
    )
    @action(detail=False, methods=["get"])
    def GetProcessListByDeviceId(self, request):
        try:
            device_id = request.query_params["device_id"]
            r = redis.Redis(host="redis-agent", port=6379, db=0)
            data = {"device_id": device_id, "process_list": []}
            json_processlist = []
            if r.get(f"process_device_{device_id}"):
                json_processlist = json.loads(r.get(f"process_device_{device_id}"))
            json_processlist = [
                x
                for x in json_processlist
                if (x.get("sha1", None) != None and x.get("sha1") != "")
            ]
            searchKey = request.query_params.get("searchKey", None)
            if r.get(f"process_device_check"):
                json_check_processlist = json.loads(r.get(f"process_device_check"))
            else:
                json_check_processlist = {}

            resList = []
            for process in json_processlist:
                threat = json_check_processlist.get(str(process.get("sha1")), None)
                process["threat"] = True
                # # if threat == None:
                # if False:
                #     virusTotalUrl = f'{os.environ.get("VIRUS_TOTAL_URL")}/{process.get("sha1")}'
                #     try:
                #         res = requests.get(virusTotalUrl, headers={"x-apikey": os.environ.get("VIRUS_TOTAL_API_KEY")})
                #     except:
                #         res = False
                #     if res and res.status_code == 200:
                #         statistic = res.json().get('data', {}).get('attributes', {}).get('last_analysis_stats', None)
                #         if statistic != None:
                #             suspicious = statistic.get('suspicious', 0)
                #             malicious = statistic.get('malicious', 0)
                #             total = suspicious + malicious
                #             if total > 0:
                #                 json_check_processlist[str(process.get('sha1'))] = True
                #                 process['threat'] = True
                #                 resList.append(process)
                #             else:
                #                 json_check_processlist[str(process.get('sha1'))] = False
                #                 process['threat'] = False
                #         else:
                #             json_check_processlist[str(process.get('sha1'))] = True
                #             process['threat'] = True
                #             resList.append(process)
                #     else:
                #         json_check_processlist[str(process.get('sha1'))] = True
                #         process['threat'] = True
                #         resList.append(process)
                # elif threat == True:
                #     resList.append(process)
                #     process['threat'] = True
                # else:
                #     process['threat'] = False

            if json_processlist and len(json_processlist) > 0:
                add_process_list_x(
                    {
                        "device_id": device_id,
                        "process_list": json_processlist,
                    }
                )
            r.set(f"process_device_check", json.dumps(json_check_processlist))
            json_processlist = resList

            if searchKey != None:
                json_processlist = [
                    x
                    for x in json_processlist
                    if (
                        (str(x.get("pid", None)) == searchKey)
                        or (searchKey in x.get("sha1"))
                    )
                ]
            filter = request.query_params.get("filter", None)
            if filter != None and str(filter).lower() == "true":
                json_processlist = [
                    x
                    for x in json_processlist
                    if (x.get("sha1", None) != None and x.get("sha1") != "")
                ]
            data["process_list"] = json_processlist
            # result = ProcessList.objects.filter(
            #     device_id=request.query_params['device_id']).order_by('-created').first()
            # serializer = self.get_serializer(result)
            return Response(data)
        except ProcessHash.DoesNotExist:
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="post",
        operation_description="post",
        # responses=openapi.Response(
        #     'ProcessListSerializer', ProcessListSerializer),
        request_body=AddProcessListWithIPSerializer,
    )
    # @api_view(['POST'])
    @action(detail=False, methods=["post"])
    def AddProcessListWithIP(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            ip_addr = xxx_data["ip"]
            current_device = Devices.objects.filter(ip=ip_addr).first()
            serializer = ProcessListSerializer(
                data={
                    **{"process_list": xxx_data["process_list"]},
                    **{"device_id": current_device.id},
                }
            )
            if serializer.is_valid(raise_exception=True):
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)


class SyscallListViewSet(viewsets.ModelViewSet):
    queryset = SyscallList.objects.all().order_by("-created")
    serializer_class = SyscallListSerializer
    permission_classes = [permissions.IsAuthenticated]

    # custom params
    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="Mã thiết bị",
        type=openapi.TYPE_INTEGER,
        required=True,
    )
    pid = openapi.Parameter(
        "pid",
        openapi.IN_QUERY,
        description="PID",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
        operation_description="get",
        # responses=openapi.Response('SyscallListSerializer', SyscallListSerializer)
    )
    @action(detail=False, methods=["get"])
    def GetSyscallListByDeviceID(self, request):
        try:
            result = SyscallList.objects.filter(
                device_id=request.query_params["device_id"]
            ).order_by("-created")
            serializer = self.get_serializer(result, many=True)
            return Response(serializer.data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id, pid],
        operation_description="get",
        # responses=openapi.Response('SyscallListSerializer', SyscallListSerializer)
    )
    @action(detail=False, methods=["get"])
    def TraceSyscall(self, request):
        try:
            currentDevice = (
                Devices.objects.filter(id=request.query_params["device_id"])
                .order_by("-created")
                .first()
            )
            # Nếu thiết bị chưa được cài agent thì hủy
            if not currentDevice.agentInstalled:
                return Response(
                    data={"message": "Thiết bị chưa được cài tác tử"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            currentDevice.tracing_syscall = request.query_params["pid"]
            currentDevice.save()
            # TODO call syscall tracing func to agent
            tasks.publish_message({"type": "stop_trace_syscall"})
            tasks.publish_message(
                {
                    "type": "install_syscall",
                    "id": currentDevice.id,
                    "ip_agent": currentDevice.ip,
                    "ip_django": os.getenv("LAN_IP"),
                    "pid": request.query_params["pid"],
                    "remote_port": currentDevice.remote_port,
                }
            )
            return Response(DevicesSerializer(currentDevice).data)

        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
        operation_description="get",
        # responses=openapi.Response('SyscallListSerializer', SyscallListSerializer)
    )
    @action(detail=False, methods=["get"])
    def StopTraceSyscall(self, request):
        try:
            currentDevice = (
                Devices.objects.filter(id=request.query_params["device_id"])
                .order_by("-created")
                .first()
            )
            currentDevice.tracing_syscall = ""
            currentDevice.save()
            # TODO call syscall tracing func to agent
            tasks.publish_message({"type": "stop_trace_syscall"})
            return Response(DevicesSerializer(currentDevice).data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
        operation_description="get",
        # responses=openapi.Response('SyscallListSerializer', SyscallListSerializer)
    )
    @action(detail=False, methods=["get"])
    def TraceNetwork(self, request):
        r = redis.Redis(host="redis-agent", port=6379, db=0)
        try:
            currentDevice = (
                Devices.objects.filter(id=request.query_params["device_id"])
                .order_by("-created")
                .first()
            )
            # Nếu thiết bị chưa được cài agent thì hủy
            if not currentDevice.agentInstalled:
                return Response(
                    data={"message": "Thiết bị chưa được cài tác tử"},
                    status=status.HTTP_200_OK,
                )
            currentDevice.tracing_network = True
            currentDevice.save()
            # TODO call pcap tracing func to agent
            data_send_check_norm = {"device": currentDevice.id}
            action_check_device_norm(data_send_check_norm, False, True)

            # register ip and device to redis
            # host_ip = os.getenv("LOCAL_HOST")
            # if r.get("device_activate_pcap"):
            #     device_activate_pcap= json.loads(r.get("device_activate_pcap"))
            #     device_list = device_activate_pcap.get(host_ip)
            #     if device_list:
            #         if currentDevice.id not in device_list:
            #             device_list.append(currentDevice.id)
            #     else: device_activate_pcap[host_ip]=[currentDevice.id]
            #     r.set("device_activate_pcap",json.dumps(device_activate_pcap))
            # else:
            #     device_activate_pcap = {host_ip:[currentDevice.id]}
            #     r.set("device_activate_pcap",json.dumps(device_activate_pcap))

            tasks.publish_message(
                {
                    "type": "stop_trace_pcap",
                    "id": currentDevice.id,
                    "ip_agent": currentDevice.ip,
                }
            )
            tasks.publish_message(
                {
                    "type": "trace_pcap",
                    "id": currentDevice.id,
                    "ip_agent": currentDevice.ip,
                    "ip_django": os.getenv("LAN_IP"),
                    "remote_port": currentDevice.remote_port,
                }
            )
            return Response(DevicesSerializer(currentDevice).data)

        except Exception as e:
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
        operation_description="get",
        # responses=openapi.Response('SyscallListSerializer', SyscallListSerializer)
    )
    @action(detail=False, methods=["get"])
    def StopTraceNetwork(self, request):
        try:
            currentDevice = (
                Devices.objects.filter(id=request.query_params["device_id"])
                .order_by("-created")
                .first()
            )
            currentDevice.tracing_network = False
            currentDevice.save()
            # TODO call pcap tracing func to agent
            data_send_check_norm = {"device": currentDevice.id}
            action_check_device_norm(data_send_check_norm, False, False)
            tasks.publish_message(
                {
                    "type": "stop_trace_pcap",
                    "id": currentDevice.id,
                    "ip_agent": currentDevice.ip,
                }
            )
            return Response(DevicesSerializer(currentDevice).data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)


class IntegrityCheckViewSet(viewsets.ModelViewSet):
    queryset = IntegrityCheck.objects.all().order_by("-created")
    serializer_class = IntegrityCheckSerializer
    permission_classes = [permissions.IsAuthenticated]

    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="device_id",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
        # operation_description="get",
        # responses=openapi.Response('IntegrityCheckSerializer', IntegrityCheckSerializer)
    )
    @action(detail=False, methods=["get"])
    def GetItegrityCheckByDeviceId(self, request):
        try:
            result = IntegrityCheck.objects.filter(
                device_id=request.query_params["device_id"]
            ).order_by("-created")
            serializer = self.get_serializer(result, many=True)
            return Response(serializer.data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="post",
        operation_description="POST",
        request_body=IntegrityCheckSerializer,
        # response=openapi.Response('IntegrityCheckSerializer', IntegrityCheckSerializer)
    )
    # @api_view(['POST'])
    @action(detail=False, methods=["post"])
    def AddInterityCheckWithIP(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            ip_addr = xxx_data["ip"]
            current_device = Devices.objects.filter(ip=ip_addr).first()
            serializer = IntegrityCheckSerializer(
                data={
                    **{"message": xxx_data["message"]},
                    **{"device_id": current_device.id},
                }
            )
            if serializer.is_valid(raise_exception=True):
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
        except Exception as e:
            logging.error(e)

            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)


class ModelMachineLearningView(
    viewsets.ModelViewSet,
    # APIView
):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = ModelMachineLearningSerializer
    queryset = ModelML.objects.all()

    @swagger_auto_schema(
        method="get",
    )
    @action(detail=False, methods=["get"])
    def check_version(self, request):
        try:
            center_server_url = os.getenv("CENTER_SERVER_URL")
            r = requests.get(center_server_url + "dlmodel/uploadModel")
            data = json.loads(r.text)
            version_input = version.parse(data["results"][0]["version"])
            queryset = (
                ModelML.objects.filter(
                    category=data["results"][0]["category"],
                    version=data["results"][0]["version"],
                )
                .order_by("-timestamp", "version")
                .count()
            )

            if queryset == 0:
                url = data["results"][0]["file"]
                r = requests.get(url)
                if url.find("/"):
                    file_name = "media/models/" + url.rsplit("/", 1)[1]
                open(file_name, "wb").write(r.content)
                data["results"][0]["file"] = re.sub(
                    center_server_url + "media/", "", data["results"][0]["file"]
                )
                serializer = ModelML.objects.create(
                    version=data["results"][0]["version"],
                    category=data["results"][0]["category"],
                    file=data["results"][0]["file"],
                )
                serializer.save()
                return Response(data={"status": True}, status=status.HTTP_201_CREATED)
            else:
                return Response(
                    data={"status": False}, status=status.HTTP_501_NOT_IMPLEMENTED
                )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)


class IpsView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = IpsTrackingSerializer
    queryset = IpsTracking.objects.all()
    parser_classes = [FormParser, MultiPartParser, JSONParser]

    ip_tracking = openapi.Parameter(
        "ip_tracking",
        openapi.IN_QUERY,
        description="IP theo doi",
        type=openapi.TYPE_STRING,
    )
    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="ID thiet bi",
        required=True,
        type=openapi.TYPE_STRING,
    )
    page = openapi.Parameter(
        "current",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_STRING,
        required=True,
    )

    page_size = openapi.Parameter(
        "page_size",
        openapi.IN_QUERY,
        description="Số lượng mỗi trang",
        type=openapi.TYPE_STRING,
        required=True,
    )
    add_count = openapi.Parameter(
        "add_count",
        openapi.IN_QUERY,
        description="So luong can them",
        type=openapi.TYPE_INTEGER,
        required=True,
    )

    @swagger_auto_schema(
        method="patch",
        request_body=UpdateIPFrequencySerializer,
    )
    @action(detail=False, methods=["patch"])
    def add_dst_ip(self, request):
        try:
            xxx_data = request.data
            for ip_track in xxx_data["ip_tracking"]:
                data = IpsTracking.objects.filter(
                    ip_tracking=ip_track,
                    device=xxx_data["device"],
                ).first()
                check = IpsTracking.objects.filter(
                    ip_tracking=ip_track,
                    device=xxx_data["device"],
                ).count()
                if check == 0:
                    ips_tracking_data = IpsTrackingSerializer(
                        data={
                            **{"device": xxx_data["device"]},
                            **{"ip_tracking": ip_track},
                            **{"count": 1},
                        }
                    )
                    if ips_tracking_data.is_valid(raise_exception=True):
                        ips_tracking_data.save()
                else:
                    pk = data.id
                    ips_tracking_data = IpsTracking.objects.get(id=pk)
                    ips_tracking_data.count = ips_tracking_data.count + 1
                    ips_tracking_data.save()
            return Response(data=True, status=status.HTTP_201_CREATED)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id, page, page_size, ip_tracking],
    )
    @action(detail=False, methods=["get"])
    def list_ip_tracking(self, request):
        try:
            result = IpsTracking.objects.filter(
                device=request.query_params["device_id"]
            ).order_by("-count")
            page_size = request.query_params["page_size"]
            page = request.query_params["current"]

            if request.query_params.get("ip_tracking", None) != None:
                result = IpsTracking.objects.filter(
                    ip_tracking=request.query_params["ip_tracking"]
                )

            paginator = Paginator(result, page_size)
            serializer = self.get_serializer(paginator.get_page(page), many=True)
            return Response(
                {
                    "results": serializer.data,
                    "count": paginator.count,
                }
            )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(method="patch", request_body=AddCountIpsTrackingSerializer)
    @action(detail=False, methods=["patch"])
    def add_count(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            result = IpsTracking.objects.filter(
                device_id=xxx_data["device"],
                ip_tracking=xxx_data["ip_tracking"],
            ).first()
            add_count = xxx_data["count"]
            result.count = result.count + int(add_count)
            result.save()
            serializer = IpsTrackingSerializer(result)
            return Response(data=serializer.data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(method="get")
    @action(detail=False, methods=["get"])
    def stat_ips(self, request):
        try:
            queryset = IpsTracking.objects.all()
            data = []
            for ip in queryset.iterator():
                result = True
                for i in range(len(data)):
                    if ip.ip_tracking == data[i][0]:
                        data[i][1] = data[i][1] + ip.count
                        result = False
                        break
                if result == True:
                    dt = []
                    dt.append(ip.ip_tracking)
                    dt.append(ip.count)
                    data.append(dt)
            total_ip = sorted(data, key=lambda x: x[1], reverse=True)
            data_json = []
            len_ip = len(total_ip) if len(total_ip) < 10 else 10
            for i in range(len_ip):
                data_json.append({"type": total_ip[i][0], "value": total_ip[i][1]})
            return Response(data=data_json)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)


class OneSignalView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser | HasAPIKey]
    serializer_class = OneSignalSerializer
    queryset = OneSignal.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [FormParser, MultiPartParser, JSONParser]
    filter_backends = [DjangoFilterBackend]
    filterset_fields = [
        "username",
        "onesignal_id",
    ]
    onesignal_id = openapi.Parameter(
        "onesignal_id",
        openapi.IN_QUERY,
        description="onesignal_id cua thiet bi di dong",
        type=openapi.TYPE_STRING,
        required=True,
    )
    username = openapi.Parameter(
        "username",
        openapi.IN_QUERY,
        description="ten tai khoan nguoi dung",
        type=openapi.TYPE_STRING,
        required=True,
    )
    notificaiton_mobile = openapi.Parameter(
        "notificaiton_mobile",
        openapi.IN_QUERY,
        description="notificaiton_mobile",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[onesignal_id, username],
    )
    @action(detail=False, methods=["get"])
    def find_device_id(self, request):
        try:
            result = OneSignal.objects.get(
                onesignal_id=request.query_params.get("onesignal_id", None),
                username=request.query_params.get("username", None),
            )
            serializer = self.get_serializer(result)
            return Response(serializer.data)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)

    # @swagger_auto_schema(
    #     method='get',
    #     manual_parameters=[notificaiton_mobile],
    # )
    # @action(detail=False, methods=['get'])
    # def push_noti_device(self, request):
    #     try:

    #         query_result_as_objects = OneSignal.objects.all() #.filter(is_superuser = True)
    #         info_device = []
    #         for e in query_result_as_objects:
    #             info_device.append(self.get_serializer(e).data['onesignal_id'])
    #         client = Client(app_id=APP_ID, rest_api_key=REST_API_KEY, user_auth_key=USER_AUTH_KEY)

    #         try:
    #             notification_body = {
    #                 'contents': {'en': request.query_params.get('notificaiton_mobile', None)},
    #                 'included_segments': ['Active Users'],
    #                 "include_external_user_ids": info_device,

    #             }

    #             # Make a request to OneSignal and parse response
    #             response = client.send_notification(notification_body)

    #         except OneSignalHTTPError as e: # An exception is raised if response.status_code != 2xx

    #         return Response(notification_body)
    #     except Exception as e:
    #         logging.error(e)
    #         return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="delete",
        manual_parameters=[onesignal_id, username],
    )
    @action(detail=False, methods=["delete"])
    def delete_device(self, request):
        try:
            result = OneSignal.objects.filter(
                onesignal_id=request.query_params.get("onesignal_id", None),
                username=request.query_params.get("username", None),
            ).first()
            result.delete()
            return Response({"status": "Delete data successfully"})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": False}, status=status.HTTP_400_BAD_REQUEST)


class RuleBaseView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = RuleBaseSerializer
    queryset = RuleBase.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [FormParser, MultiPartParser, JSONParser]
    filter_backends = [DjangoFilterBackend]
    rules = openapi.Parameter(
        "rules",
        openapi.IN_QUERY,
        description="Rules",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[rules],
    )
    @action(detail=False, methods=["get"])
    def rulebase_info(self, request):
        if request.query_params.get("rules", None) == "firewall":
            output = (
                RuleBase.objects.filter(rules="firewall").order_by("-timestamp").first()
            )
        else:
            output = (
                RuleBase.objects.filter(rules="snort").order_by("-timestamp").first()
            )
        return Response(RuleBaseSerializer(output).data)


class DeviceIpConnectHWagentView(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = DeviceIpConnectHWagentSerializer
    queryset = DeviceIpConnectHWagent.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [FormParser, MultiPartParser, JSONParser]
    filter_backends = [DjangoFilterBackend]

    def get_object(self, pk):
        try:
            return Devices.objects.get(pk=pk)
        except Devices.DoesNotExist:
            raise Http404

    data_hw = openapi.Parameter(
        "data_hw",
        openapi.IN_QUERY,
        description="data_hw",
        type=openapi.TYPE_STRING,
        required=True,
    )
    # @swagger_auto_schema(
    #     method='patch',
    #     request_body=ListDeviceIpConnectHWagentSerializer
    # )
    # @action(detail=False, methods=['patch'])
    # def update_device_connect_hwagent(self, request):
    #     try:
    #         xxx_data = json.loads(request.query_params['data'])
    #         result = DeviceIpConnectHWagent.objects.filter(device = xxx_data['device'])
    #         count = DeviceIpConnectHWagent.objects.filter(device = xxx_data['device']).count()
    #         leg_data = len(xxx_data['ip'])
    #         if count > 0:
    #             for dbsave in result.iterator():
    #                 dbsave.connect = 'N'
    #                 dbsave.save()

    #             for stt in range(leg_data):
    #                 res = False
    #                 for dbsave in result.iterator():
    #                         if xxx_data['mac_addr'][stt] == dbsave.mac_addr:
    #                             dbsave.connect = 'D'
    #                             res = True
    #                             dbsave.save()
    #                 if res == False:
    #                     dbsave = DeviceIpConnectHWagent(
    #                         device = self.get_object(pk = xxx_data['device']),
    #                         mac_addr = xxx_data['mac_addr'][stt],
    #                         ip = xxx_data['ip'][stt],
    #                         hostname=xxx_data['hostname'][stt]
    #                     )
    #                     dbsave.save()
    #         else:
    #             for stt in range(leg_data):
    #                     dbsave = DeviceIpConnectHWagent(
    #                         device = self.get_object(pk = xxx_data['device']),
    #                         mac_addr = xxx_data['mac_addr'][stt],
    #                         ip = xxx_data['ip'][stt],
    #                         hostname=xxx_data['hostname'][stt]
    #                     )
    #                     dbsave.save()
    #         return Response({"status": True})
    #     except Exception as e:
    #         logging.error(e)
    #         return Response(data={'status': e}, status=status.HTTP_400_BAD_REQUEST)

    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="device_id",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
    )
    @action(detail=False, methods=["get"])
    def get_ipdevice(self, request):
        try:
            device = self.get_object(pk=request.query_params.get("device_id", None))
            result = DeviceIpConnectHWagent.objects.filter(
                device=device.id, network=False
            )
            data = []
            for dt in result.iterator():
                data.append(dt.ip)
            return Response({"data": data})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": e}, status=status.HTTP_400_BAD_REQUEST)

    connect = openapi.Parameter(
        "connect",
        openapi.IN_QUERY,
        description="connect",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id, connect],
    )
    @action(detail=False, methods=["get"])
    def get_ip_in_device(self, request):
        try:
            device = self.get_object(pk=request.query_params.get("device_id", None))
            result = []
            if request.query_params.get("connect", None) == "D":
                result = DeviceIpConnectHWagent.objects.filter(
                    device=device.id, connect="D"
                )
            else:
                result = DeviceIpConnectHWagent.objects.filter(device=device.id)
            data = serializers.serialize("json", result)
            return Response(json.loads(data))
        except Exception as e:
            logging.error(e)
            return Response(data={"status": e}, status=status.HTTP_400_BAD_REQUEST)

    network = openapi.Parameter(
        "network",
        openapi.IN_QUERY,
        description="network",
        type=openapi.TYPE_STRING,
        required=True,
    )

    @swagger_auto_schema(
        method="patch",
        request_body=LisstDeviceIdConnectHWagentSerializer,
        manual_parameters=[network, device_id],
    )
    @action(detail=False, methods=["patch"])
    def update_device_network_hwagent(self, request):
        try:
            xxx_data = json.loads(request.query_params["data"])
            for dt in xxx_data["id"]:
                result = DeviceIpConnectHWagent.objects.get(pk=dt)
                result.network = request.query_params.get("network", None)
                result.save()
            current_device = Devices.objects.filter(
                id=request.query_params.get("device_id", None)
            ).first()
            if current_device.hwagent:
                tasks.publish_message(
                    {
                        "type": "update rulebase hw",
                        "id": current_device.id,
                        "ip_agent": current_device.ip,
                        "protocol": current_device.protocol,
                        "ip_django": os.getenv("LAN_IP"),
                        "ip_sock_serv": os.getenv("LAN_SOCKET_SERVER_IP")
                        if not current_device.from_internet
                        else os.getenv("PUBLIC_SOCKET_SERVER"),
                        "ip_tftp_serv": os.getenv("LAN_SOCKET_SERVER_IP")
                        if not current_device.from_internet
                        else os.getenv("PUBLIC_TFTP_SERVER"),
                        "username": current_device.username,
                        "password": current_device.password,
                        "OS": current_device.os,
                        "port": current_device.port,
                        "remote_port": current_device.remote_port,
                        "hwagent": current_device.hwagent,
                    }
                )
            return Response({"status": True})
        except Exception as e:
            logging.error(e)
            return Response(data={"status": e}, status=status.HTTP_400_BAD_REQUEST)


class CreateDevice(generics.CreateAPIView):
    serializer_class = DevicesSerializer
    permission_classes = [permissions.IsAdminUser]
    parser_classes = (FormParser, MultiPartParser)

    def post(self, request, *args, **kwargs):
        # filename = xxx_data['file']
        return self.create(request, *args, **kwargs)


class EditDevice(generics.CreateAPIView):
    serializer_class = DevicesSerializer
    permission_classes = [permissions.IsAdminUser]
    parser_classes = (FormParser, MultiPartParser)

    def get_object(self, pk):
        try:
            return Devices.objects.get(pk=pk)
        except Devices.DoesNotExist:
            raise Http404

    def patch(self, request, *args, **kwargs):
        try:
            xxx_data = json.loads(request.query_params["data"])
            device = self.get_object(ip=xxx_data["ip"], port=xxx_data["port"])
            device.avatar = request.FILES["avatars"]
            device.save()

            return True
        except Exception as e:
            logging.error(e)
            return Response(data={"status": e}, status=status.HTTP_400_BAD_REQUEST)


class AutoUpdateListDevice(APIView):
    subnet = openapi.Parameter(
        "subnet", openapi.IN_QUERY, description="subnet", type=openapi.TYPE_STRING
    )

    @swagger_auto_schema(
        manual_parameters=[subnet],
    )
    def post(self, request):
        config = dotenv_values("/backend/.env.dev")

        subnet = request.query_params["subnet"]
        ip_host = config.get("LAN_SOCKET_SERVER_IP")
        username = config.get("HOST_USERNAME")
        password = config.get("HOST_PASSWORD")
        interface = config.get("HOST_INTERFACE")
        ssh = SshClient()
        login = ssh.login_host_by_ssh(ip_host, username, password, 22)
        if login == True:
            scan_device = ssh.scan_mac_device(subnet, interface, password)
        else:
            return False

        return HttpResponse("Changed seconds auto update snort.")


class ThresholdView(APIView):
    def get(self, request):
        try:
            # config = dotenv_values("/backend/.env.dev")
            cpu_threshold = int(get_env("CPU_THRESHOLD"))
            ram_threshold = int(get_env("RAM_THRESHOLD"))
            latency_threshold = int(get_env("LATENCY_THRESHOLD"))
            active_distributed = get_env("ACTIVE_DISTRIBUTED").lower() == "true"
            active_distributed_receive = (
                get_env("ACTIVE_DISTRIBUTED_RECEIVE").lower() == "true"
            )
            return Response(
                {
                    "cpu_threshold": cpu_threshold,
                    "ram_threshold": ram_threshold,
                    "latency_threshold": latency_threshold,
                    "active_distributed": active_distributed,
                    "active_distributed_receive": active_distributed_receive,
                },
                status=status.HTTP_200_OK,
            )
        except Exception as e:
            return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    cpu_threshold = openapi.Parameter(
        name="cpu_threshold",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_INTEGER,
        required=True,
    )
    ram_threshold = openapi.Parameter(
        name="ram_threshold",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_INTEGER,
        required=True,
    )
    latency_threshold = openapi.Parameter(
        name="latency_threshold",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_INTEGER,
        required=True,
    )
    active_distributed = openapi.Parameter(
        name="active_distributed",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_BOOLEAN,
        required=True,
    )
    active_distributed_receive = openapi.Parameter(
        name="active_distributed_receive",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_BOOLEAN,
        required=True,
    )

    threshold_schema = openapi.Schema(
        type=openapi.TYPE_OBJECT,
        properties={
            "cpu_threshold": cpu_threshold,
            "ram_threshold": ram_threshold,
            "latency_threshold": latency_threshold,
            "active_distributed": active_distributed,
            "active_distributed_receive": active_distributed_receive,
        },
    )

    @swagger_auto_schema(
        request_body=threshold_schema,
    )
    def put(self, request):
        try:
            data = request.data
            contents = Path(pathEnv).read_text()
            arr = contents.split(chr(10))
            update_env("CPU_THRESHOLD", str(data.get("cpu_threshold", 90)))
            update_env("RAM_THRESHOLD", str(data.get("ram_threshold", 90)))
            update_env("LATENCY_THRESHOLD", str(data.get("latency_threshold", 10)))
            update_env(
                "ACTIVE_DISTRIBUTED_RECEIVE",
                str(data.get("active_distributed_receive", True)),
            )
            update_env("ACTIVE_DISTRIBUTED", str(data.get("active_distributed", True)))
            for i in range(len(arr)):
                if arr[i].find("CPU_THRESHOLD") != -1:
                    arr[i] = "CPU_THRESHOLD=" + str(data.get("cpu_threshold", 90))
                if arr[i].find("RAM_THRESHOLD") != -1:
                    arr[i] = "RAM_THRESHOLD=" + str(data.get("ram_threshold", 90))
                if arr[i].find("LATENCY_THRESHOLD") != -1:
                    arr[i] = "LATENCY_THRESHOLD=" + str(
                        data.get("latency_threshold", 10)
                    )
                if arr[i].find("ACTIVE_DISTRIBUTED_RECEIVE") != -1:
                    arr[i] = "ACTIVE_DISTRIBUTED_RECEIVE=" + str(
                        data.get("active_distributed_receive", True)
                    )
                elif arr[i].find("ACTIVE_DISTRIBUTED") != -1:
                    arr[i] = "ACTIVE_DISTRIBUTED=" + str(
                        data.get("active_distributed", True)
                    )
            res = ""
            for i in range(len(arr)):
                res += arr[i]
                if i < len(arr) - 1:
                    res += chr(10)
            open(pathEnv, "w").close()
            f = open(pathEnv, "w")
            f.write(res)
            f.close()
            return Response("Changed threshold", status=status.HTTP_200_OK)
        except Exception as e:
            return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class GetExportPDFURL(APIView):
    def get(self, request):
        url = request.get_full_path()
        url = url[: url.find("/url")] + url[url.find("/url") + 4 :]
        return Response(url)


class ExportPDF(APIView):
    def get(self, request):
        try:
            xxx_data = json.loads(request.query_params.get("data", None))
            result = Alerts.objects.all()
            if request.query_params.get("ip", None) != None:
                result = result.filter(ip__icontains=request.query_params["ip"])
            if request.query_params.get("message", None) != None:
                result = result.filter(
                    message__icontains=request.query_params["message"]
                )
            if request.query_params.get("address", None) != None:
                result = result.filter(
                    address__icontains=request.query_params["address"]
                )

            if xxx_data.get("timestamp", None) != None:
                result = result.filter(
                    timestamp__range=(
                        xxx_data.get("timestamp", None)[0],
                        xxx_data.get("timestamp", None)[1],
                    )
                )
            i = 0
            xyz = result
            if request.query_params["all_device"] == "false":
                if xxx_data.get("device", None) != None:
                    for id in xxx_data.get("device", None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0

            if xxx_data.get("type", None) != None:
                for id in xxx_data.get("type", None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = abc | result.filter(type=id)
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get("status", None) != None:
                for id in xxx_data.get("status", None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz

            buffer = io.BytesIO()
            document = Document()

            document.add_heading("Cảnh báo tấn công mạng", 0)

            table = document.add_table(rows=1, cols=8)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "STT"
            hdr_cells[1].text = "Mã thiết bị"
            hdr_cells[2].text = "Địa chỉ ip"
            hdr_cells[3].text = "Cảnh báo"
            hdr_cells[4].text = "Hash"
            hdr_cells[5].text = "PID"
            hdr_cells[6].text = "Thời gian"
            hdr_cells[7].text = "Địa chỉ"

            for i in range(len(xyz)):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = str(xyz[i].device_id)
                row_cells[2].text = str(xyz[i].ip)
                row_cells[3].text = str(xyz[i].message)
                row_cells[4].text = str(xyz[i].hash)
                row_cells[5].text = str(xyz[i].pid)
                row_cells[6].text = str(xyz[i].timestamp.strftime("%d/%m/%Y, %H:%M:%S"))
                row_cells[7].text = str(xyz[i].address)

            table.style = "Colorful List"

            document.save("./devices/reports/alerts.docx")
            convert_to_pdf = f"libreoffice --headless --convert-to pdf ./devices/reports/alerts.docx --outdir ./devices/reports/pdf"
            subprocess.run(convert_to_pdf, shell=True)
            return FileResponse(
                open("./devices/reports/pdf/alerts.pdf", "rb+"),
                as_attachment=True,
                filename="alerts.pdf",
                status=status.HTTP_200_OK,
            )
        except Exception as e:
            logging.error(e)
            return Response(
                data={"status": False, "Log bug": str(e)},
                status=status.HTTP_400_BAD_REQUEST,
            )


class GetExportDOCXURL(APIView):
    def get(self, request):
        url = request.get_full_path()
        url = url[: url.find("/url")] + url[url.find("/url") + 4 :]
        return Response(url)


class ExportDOCX(APIView):
    def get(self, request):
        try:
            xxx_data = json.loads(request.query_params.get("data", None))
            result = Alerts.objects.all()
            if request.query_params.get("ip", None) != None:
                result = result.filter(ip__icontains=request.query_params["ip"])
            if request.query_params.get("message", None) != None:
                result = result.filter(
                    message__icontains=request.query_params["message"]
                )
            if request.query_params.get("address", None) != None:
                result = result.filter(
                    address__icontains=request.query_params["address"]
                )

            if xxx_data.get("timestamp", None) != None:
                result = result.filter(
                    timestamp__range=(
                        xxx_data.get("timestamp", None)[0],
                        xxx_data.get("timestamp", None)[1],
                    )
                )
            i = 0
            xyz = result
            if request.query_params["all_device"] == "false":
                if xxx_data.get("device", None) != None:
                    for id in xxx_data.get("device", None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0

            if xxx_data.get("type", None) != None:
                for id in xxx_data.get("type", None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = abc | result.filter(type=id)
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get("status", None) != None:
                for id in xxx_data.get("status", None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz

            # buffer = io.BytesIO()
            document = Document()

            document.add_heading("Cảnh báo tấn công mạng", 0)

            table = document.add_table(rows=1, cols=8)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "STT"
            hdr_cells[1].text = "Mã thiết bị"
            hdr_cells[2].text = "Địa chỉ ip"
            hdr_cells[3].text = "Cảnh báo"
            hdr_cells[4].text = "Hash"
            hdr_cells[5].text = "PID"
            hdr_cells[6].text = "Thời gian"
            hdr_cells[7].text = "Địa chỉ"

            for i in range(len(xyz)):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = str(xyz[i].device_id)
                row_cells[2].text = str(xyz[i].ip)
                row_cells[3].text = str(xyz[i].message)
                row_cells[4].text = str(xyz[i].hash)
                row_cells[5].text = str(xyz[i].pid)
                row_cells[6].text = str(xyz[i].timestamp.strftime("%d/%m/%Y, %H:%M:%S"))
                row_cells[7].text = str(xyz[i].address)

            table.style = "Colorful List"

            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response["Content-Disposition"] = "attachment; filename=alerts.docx"
            # response['Content-Length'] = buffer.tell()
            document.add_page_break()
            document.save(response)
            # document.save("./devices/reports/alerts.docx")
            return response
        except Exception as e:
            logging.error(e)
            return Response(
                data={"status": False, "Log bug": str(e)},
                status=status.HTTP_400_BAD_REQUEST,
            )


class BlacklistClearView(APIView):
    def get(self, request):
        try:
            BlackListIP.objects.all().delete()
            return Response(
                {"message": "Clear blacklist successfull"}, status=status.HTTP_200_OK
            )
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class WhitelistClearView(APIView):
    def get(self, request):
        try:
            WhiteListIP.objects.all().delete()
            return Response(
                {"message": "Clear whitelist successfull"}, status=status.HTTP_200_OK
            )
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class AgentClearFileSFTPView(APIView):
    def get(self, request):
        try:
            return Response(
                int(r_agent.get("AUTO_CLEAR_SFTP_STORAGE_MINUTES")),
                status=status.HTTP_200_OK,
            )
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    minutes = openapi.Parameter(
        name="minutes",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_INTEGER,
        required=True,
    )

    agent_sftp_auto_clear_schema = openapi.Schema(
        type=openapi.TYPE_OBJECT,
        properties={
            "minutes": minutes,
        },
    )

    @swagger_auto_schema(
        request_body=agent_sftp_auto_clear_schema,
    )
    def put(self, request):
        try:
            data = request.data

            r_agent.set("AUTO_CLEAR_SFTP_STORAGE_MINUTES", data.get("minutes", 30))
            return Response(
                {"message": "Update AUTO_CLEAR_SFTP_STORAGE_MINUTES successfull"},
                status=status.HTTP_200_OK,
            )
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class AgentCheckOverloadView(APIView):
    def get(self, request):
        try:
            return Response(
                int(r_agent.get("AUTO_CHECK_OVERLOAD_SECONDS")),
                status=status.HTTP_200_OK,
            )
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    seconds = openapi.Parameter(
        name="seconds",
        in_=openapi.IN_BODY,
        type=openapi.TYPE_INTEGER,
        required=True,
    )

    agent_auto_check_overload_schema = openapi.Schema(
        type=openapi.TYPE_OBJECT,
        properties={
            "seconds": seconds,
        },
    )

    @swagger_auto_schema(
        request_body=agent_auto_check_overload_schema,
    )
    def put(self, request):
        try:
            data = request.data

            r_agent.set("AUTO_CHECK_OVERLOAD_SECONDS", data.get("seconds", 30))
            return Response(
                {"message": "Update AUTO_CHECK_OVERLOAD_SECONDS successfull"},
                status=status.HTTP_200_OK,
            )
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class DeviceNormViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAdminUser]
    serializer_class = DeviceNormSerializer
    queryset = DeviceNorm.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend]

    device = openapi.Parameter(
        "device",
        openapi.IN_QUERY,
        description="ID thiet bi",
        type=openapi.TYPE_STRING,
        required=True,
    )
    device_id = openapi.Parameter(
        "device_id",
        openapi.IN_QUERY,
        description="ID thiet bi",
        type=openapi.TYPE_INTEGER,
        required=True,
    )

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        action_device_norm("create", serializer.data)
        action_check_device_norm(serializer.data)
        headers = self.get_success_headers(serializer.data)
        return Response(
            serializer.data, status=status.HTTP_201_CREATED, headers=headers
        )

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop("partial", False)
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        self.perform_update(serializer)
        action_device_norm("update", serializer.data)
        return Response(serializer.data)

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = DeviceNormSerializer(instance)
        action_device_norm("delete", serializer.data)
        self.perform_destroy(instance)
        return Response(status=status.HTTP_204_NO_CONTENT)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device],
    )
    @action(detail=False, methods=["get"])
    def pass_device_norm(self, request):
        try:
            data = {"device": request.query_params["device"]}
            resp = action_check_device_norm(data, True)
            return Response(data=resp, status=status.HTTP_200_OK)
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="get",
        manual_parameters=[device_id],
    )
    @action(detail=False, methods=["get"])
    def get_by_device_id(self, request, *args, **kwargs):
        try:
            device_id = request.query_params.get("device_id")
            res = self.queryset.filter(device_id=device_id).first()
            if res != None:
                return Response(
                    data=DeviceNormSerializer(res).data, status=status.HTTP_200_OK
                )
            else:
                data_device_norm = {"device": device_id}
                serializer_device_norm = DeviceNormSerializer(data=data_device_norm)
                if serializer_device_norm.is_valid():
                    serializer_device_norm.save()
                    action_device_norm("create", serializer_device_norm.data)
                    action_check_device_norm(serializer_device_norm.data)
                return Response(
                    serializer_device_norm.data, status=status.HTTP_201_CREATED
                )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method="patch",
        request_body=DeviceNormSerializer,
        manual_parameters=[device_id],
    )
    @action(detail=False, methods=["patch"])
    def update_by_device_id(self, request, *args, **kwargs):
        try:
            device_id = request.query_params.get("device_id")
            data = request.data
            instance = self.queryset.filter(device_id=device_id).first()
            if instance != None:
                serializer = DeviceNormSerializer(instance, data=data)
                if serializer.is_valid():
                    serializer.save()
                    action_device_norm("update", serializer.data)
                else:
                    return Response(
                        str(serializer.errors), status=status.HTTP_400_BAD_REQUEST
                    )
                return Response(data=serializer.data, status=status.HTTP_200_OK)
            else:
                return Response(
                    "DeviceNorm Not found", status=status.HTTP_404_NOT_FOUND
                )
        except Exception as e:
            logging.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class GetFileFromPcap(APIView):
    # path = openapi.Parameter(
    #     name = "path",
    #     in_=openapi.IN_BODY,
    #     type=openapi.TYPE_STRING,
    #     required=True,
    # )
    # @swagger_auto_schema(
    #     request_body=path,
    # )
    def get(self, request):
        try:
            data = request.data
            file_path = data["path"]
            file = file_path.split("/")[-1]
            f = open(file_path)
            response = HttpResponse(f)
            response["Content-Disposition"] = f"attachment; filename={file}"
            return response
        except Exception as e:
            return Response(
                {"message": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def post(self, request):
        try:
            try:
                directory_path = "/backend/media/file_save_from_pcap"
                files = os.listdir(directory_path)
                for file in files:
                    file_path = os.path.join(directory_path, file)
                    if os.path.isfile(file_path):
                        os.remove(file_path)
            except:
                pass
            file_uploaded = request.FILES.get("file")
            path = default_storage.save(
                file_uploaded.name, ContentFile(file_uploaded.read())
            )
            tmp_file = os.path.join(settings.MEDIA_ROOT, path)
            extract_allFolder(tmp_file)
            data_response = []
            try:
                directory_path = "media/file_save_from_pcap"
                files = os.listdir(directory_path)
                for file in files:
                    file_path = os.path.join(directory_path, file)
                    file_tmp = file.replace("%20", "_")
                    file_path_tmp = os.path.join(directory_path, file_tmp)
                    os.rename(file_path, file_path_tmp)
                    data_response.append({"name": file_tmp, "path": file_path_tmp})
            except Exception as e:
                return Response(
                    {"message": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            return Response(data_response, status=status.HTTP_200_OK)
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class GetFileFromNetwork(APIView):
    def get():
        try:
            try:
                directory_path = "/backend/media/file_save_from_pcap"
                files = os.listdir(directory_path)
                for file in files:
                    file_path = os.path.join(directory_path, file)
                    if os.path.isfile(file_path):
                        os.remove(file_path)
            except:
                pass
            file_uploaded = request.FILES.get("file")
            path = default_storage.save(
                file_uploaded.name, ContentFile(file_uploaded.read())
            )
            tmp_file = os.path.join(settings.MEDIA_ROOT, path)
            extract_allFolder(tmp_file)
            data_response = []
            try:
                directory_path = "media/file_save_from_pcap"
                files = os.listdir(directory_path)
                for file in files:
                    file_path = os.path.join(directory_path, file)
                    file_tmp = file.replace("%20", "_")
                    file_path_tmp = os.path.join(directory_path, file_tmp)
                    os.rename(file_path, file_path_tmp)
                    data_response.append({"name": file_tmp, "path": file_path_tmp})
            except Exception as e:
                return Response(
                    {"message": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            return Response(data_response, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {"message": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class GetPcapJsonFromElasticsearch(viewsets.ModelViewSet):
    # permission_classes = [permissions.IsAdminUser]
    serializer_class = BlackListIPallSerializer
    queryset = BlackListIP.objects.all()
    # permission_classes = [permissions.IsAuthenticated]

    page_size = openapi.Parameter(
        "page_size",
        openapi.IN_QUERY,
        description="Số lượng mỗi trang",
        type=openapi.TYPE_INTEGER,
        required=True,
    )
    page = openapi.Parameter(
        "page",
        openapi.IN_QUERY,
        description="Trang",
        type=openapi.TYPE_INTEGER,
        required=True,
    )

    @swagger_auto_schema(
        method="get",
        manual_parameters=[page, page_size],
    )
    @action(detail=False, methods=["get"])
    def get_data_pcap(self, request):
        try:
            page_size = request.query_params["page_size"]
            page = request.query_params["page"]
            if int(page) < 1:
                page = 1
            if int(page_size) < 1:
                page_size = 1
            es = Elasticsearch(os.environ.get("ELASTIC_HOST"))
            resp = es.search(
                index="pcap_save",
                query={"match_all": {}},
                size=page_size,
                from_=int(page_size) * (int(page) - 1),
            )
            data_response = resp["hits"]["hits"]
            return Response(data_response, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {"message": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class GetFileFromPcapWithTime(APIView):
    def post(self, request):
        try:
            try:
                directory_path = "/backend/media/file_save_from_pcap"
                files = os.listdir(directory_path)
                for file in files:
                    file_path = os.path.join(directory_path, file)
                    if os.path.isfile(file_path):
                        os.remove(file_path)
            except:
                pass
            data_request = request.data
            device_id = data_request.get("device_id")
            time_start = data_request.get("start")
            time_end = data_request.get("end")
            key_name_file = f"network_{device_id}_"
            command = f'ls pcap_save | grep "{key_name_file}"'
            datas = os.popen(command).readlines()
            command_mergecap = f"mergecap -w /backend/pcap_save/merge_cap_tmp.pcap "
            # arr_file_name = []
            check_exit_file_pcap = False
            for data in datas:
                if (
                    data[:-1] > f"{key_name_file}{time_start}.pcap"
                    and data[:-1] < f"{key_name_file}{time_end}.pcap"
                ):
                    command_mergecap += f"/backend/pcap_save/{data[:-1]} "
                    check_exit_file_pcap = True
                    # arr_file_name.append(data[:-1])
            if check_exit_file_pcap:
                os.popen(command_mergecap).readlines()
                extract_allFolder("/backend/pcap_save/merge_cap_tmp.pcap")
                data_response = []
                try:
                    directory_path = "media/file_save_from_pcap"
                    files = os.listdir(directory_path)
                    for file in files:
                        file_path = os.path.join(directory_path, file)
                        file_tmp = file.replace("%20", "_")
                        file_path_tmp = os.path.join(directory_path, file_tmp)
                        os.rename(file_path, file_path_tmp)
                        data_response.append({"name": file_tmp, "path": file_path_tmp})
                except Exception as e:
                    return Response(
                        {"message": str(e)},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    )
                return Response(data_response, status=status.HTTP_200_OK)
            else:
                return Response([], status=status.HTTP_200_OK)
        except Exception as err:
            return Response(
                {"message": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class DashboardView(APIView):
    serializer_class = DevicesSerializer
    queryset = Devices.objects.all()

    def get(self, request):
        try:
            used_ram = round(psutil.virtual_memory()[2] / 100, 4)
            used_cpu = round(psutil.cpu_percent() / 100, 4)
            trace_pcap = get_env("TRACE_PCAP", "false") == "true"
            # Send data to elastic
            return Response(
                data={
                    "device_counts": 0,
                    "monitored_devices_counts": 0,
                    "free_devices_counts": 0,
                    "alert_counts": 0,
                    "reviewed_alert_counts": 0,
                    "pending_alert_counts": 0,
                    "malware_alert_counts": 0,
                    "syscall_alert_counts": 0,
                    "network_alert_counts": 0,
                    "RAM": used_ram,
                    "CPU": used_cpu,
                    "TRACE_PCAP": trace_pcap,
                },
                status=200,
            )

        except Exception as e:
            logger.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class DashboardStartTracePcapView(APIView):
    serializer_class = DevicesSerializer
    queryset = Devices.objects.all()

    def get(self, request):
        try:
            used_ram = round(psutil.virtual_memory()[2] / 100, 4)
            used_cpu = round(psutil.cpu_percent() / 100, 4)
            update_env("TRACE_PCAP", "true")
            trace_pcap = get_env("TRACE_PCAP", "false") == "true"

            threading.Thread(
                target=capture.cap,
                args=(
                    get_env("ANALYZER_INF", "wlp7s0"), 
                    1000
                ),
            ).start()
            # threading.Thread(target=capture.cap_2).start()
            return Response(
                data={
                    "device_counts": 0,
                    "monitored_devices_counts": 0,
                    "free_devices_counts": 0,
                    "alert_counts": 0,
                    "reviewed_alert_counts": 0,
                    "pending_alert_counts": 0,
                    "malware_alert_counts": 0,
                    "syscall_alert_counts": 0,
                    "network_alert_counts": 0,
                    "RAM": used_ram,
                    "CPU": used_cpu,
                    "TRACE_PCAP": trace_pcap,
                },
                status=200,
            )
        except Exception as e:
            logger.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class DashboardStopTracePcapView(APIView):
    serializer_class = DevicesSerializer
    queryset = Devices.objects.all()

    def get(self, request):
        try:
            used_ram = round(psutil.virtual_memory()[2] / 100, 4)
            used_cpu = round(psutil.cpu_percent() / 100, 4)
            update_env("TRACE_PCAP", "false")
            trace_pcap = get_env("TRACE_PCAP", "false") == "true"
            return Response(
                data={
                    "device_counts": 0,
                    "monitored_devices_counts": 0,
                    "free_devices_counts": 0,
                    "alert_counts": 0,
                    "reviewed_alert_counts": 0,
                    "pending_alert_counts": 0,
                    "malware_alert_counts": 0,
                    "syscall_alert_counts": 0,
                    "network_alert_counts": 0,
                    "RAM": used_ram,
                    "CPU": used_cpu,
                    "TRACE_PCAP": trace_pcap,
                },
                status=200,
            )
        except Exception as e:
            logger.error(e)
            return Response(data={"status": str(e)}, status=status.HTTP_400_BAD_REQUEST)


def send_cpu_ram_to_local_elastic():
    try:
        used_ram = round(psutil.virtual_memory()[2] / 100, 4)
        used_cpu = round(psutil.cpu_percent() / 100, 4)
        es = Elasticsearch(get_env("ANALYZER_ELASTIC", "http://127.0.0.1:9200"))
        delta_7h = timedelta(hours=7)
        phan_tai = get_env("PHAN_TAI", "false") == "true"
        doc = {
            "cpu": used_cpu * 100,
            "ram": used_ram * 100,
            "timestamp": datetime.now() - delta_7h,
        }
        es_index = get_env("ANALYZER_ELASTIC_INDEX", "cpu_ram")
        if not es.indices.exists(index=es_index):
            es.indices.create(index=es_index)
        es.index(index=es_index, document=doc)
    except Exception as e:
        logger.error(e)


def send_cpu_ram_to_local_redis():
    try:
        used_ram = round(psutil.virtual_memory()[2] / 100, 4)
        used_cpu = round(psutil.cpu_percent() / 100, 4)
        r = redis.Redis(
            host=get_env("ANALYZER_REDIS_HOST", "127.0.0.1"),
            port=int(get_env("ANALYZER_REDIS_PORT", 6379)),
            db=0,
        )
        phan_tai = get_env("PHAN_TAI", "false") == "true"
        r.set("cpu", used_cpu * 100)
        r.set("ram", used_ram * 100)
    except Exception as e:
        logger.error(e)


def send_cpu_ram_to_center_elastic():
    try:
        used_ram = round(psutil.virtual_memory()[2] / 100, 4)
        used_cpu = round(psutil.cpu_percent() / 100, 4)
        es = Elasticsearch(get_env("CENTER_ELASTIC", "http://127.0.0.1:9200"))
        delta_7h = timedelta(hours=7)
        phan_tai = get_env("PHAN_TAI", "false") == "true"
        doc = {
            "cpu": used_cpu * 100,
            "ram": used_ram * 100,
            "analyzer_ip": get_env("ANALYZER_IP_LOCAL", "127.0.0.1"),
            "timestamp": datetime.now() - delta_7h,
        }
        es_index = get_env("CENTER_ELASTIC_INDEX", "cpu_ram")
        if not es.indices.exists(index=es_index):
            es.indices.create(index=es_index)
        es.index(index=es_index, document=doc)
    except Exception as e:
        logger.error(e)


def send_cpu_ram_to_center_redis():
    try:
        used_ram = round(psutil.virtual_memory()[2] / 100, 4)
        used_cpu = round(psutil.cpu_percent() / 100, 4)
        r = redis.Redis(
            host=get_env("CENTER_REDIS_HOST", "127.0.0.1"),
            port=int(get_env("CENTER_REDIS_PORT", 6379)),
            db=0,
        )
        phan_tai = get_env("PHAN_TAI", "false") == "true"
        doc = {
            "cpu": used_cpu * 100,
            "ram": used_ram * 100,
        }
        redis_index = "cpu_ram_" + get_env("ANALYZER_IP_LOCAL", "127.0.0.1")
        r.set(redis_index, json.dumps(doc))
    except Exception as e:
        logger.error(e)


def get_best_analyzer():
    try:
        req = requests.get(url=f'{os.getenv("CENTER_HOST")}/devices/best-analyzer')
        if req.status_code == 200:
            datas = req.json()
            device_lefts = []
            sum_weights = 0
            weights = []
            for data in datas:
                if data.get("ip") != get_env("ANALYZER_IP_LOCAL"):
                    sum_weights += int(data.get("weight"))
                    device_lefts.append(
                        {"ip": data.get("ip"), "weight": int(data.get("weight"))}
                    )
            for device in device_lefts:
                weights.append(
                    {
                        "ip": device.get("ip"),
                        "weight": round(
                            ((int(device.get("weight"))) / sum_weights) * 10
                        ),
                    }
                )
            weights = sorted(weights, reverse=True, key=lambda d: d["weight"])
            r.set("list_analyzer", json.dumps(weights))
            print(json.dumps(weights))
    except Exception as e:
        logger.error(e)

def check_phan_tai():
    try:
        used_ram = round(psutil.virtual_memory()[2] / 100, 4)
        used_cpu = round(psutil.cpu_percent() / 100, 4)
        phan_tai = get_env("PHAN_TAI", "false") == "true"
        # if (used_cpu > 80 or used_ram > 90) and phan_tai == False:
        if used_cpu > 0.8:
            if phan_tai == False:
                print("======================SAP PHAN TAI======================", used_ram, used_cpu, phan_tai)
                update_env("PHAN_TAI", "true")
                time.sleep(30)
                update_env("PHAN_TAI_NGUONG", "true")
            else:
                print("======================SAP TIEP TUC PHAN TAI======================", used_ram, used_cpu, phan_tai)
        else:
            print("======================VAN SE TU XU LY======================", used_ram, used_cpu, phan_tai)
            update_env("PHAN_TAI", "false")
            time.sleep(30)
            update_env("PHAN_TAI_NGUONG", "false")
            # processes = []
            # for _ in range(9):
            #     p = multiprocessing.Process(target=capture.cpu_bound_task_2, args=(10000000,))
            #     p.start()
            #     processes.append(p)

            # for p in processes:
            #     p.join()
    except Exception as e:
        logger.error(e)
